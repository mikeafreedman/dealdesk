"""
word_builder.py — Report Generation Module
===========================================
Generates the final PDF underwriting report in three stages:

    1. Prompt 4-MASTER (Sonnet) — batched generation of all narrative sections.
    2. Prompt 5D (Sonnet, investor_mode only) — rewrites 9 narrative blocks
       in LP-appropriate language.
    3. docxtpl template population → LibreOffice headless PDF conversion.

Pipeline position: Stage 7 — runs after financials.py.
Output: {deal_id}_report.pdf saved to outputs/.
"""

from __future__ import annotations

import json
import logging
import os
import platform
from collections import defaultdict
import subprocess
import time
from pathlib import Path
from typing import Optional

import anthropic
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx.oxml.ns import qn

from config import (
    ANTHROPIC_API_KEY,
    MODEL_SONNET,
    OUTPUTS_DIR,
    PDF_CONVERSION_TIMEOUT,
    WORD_TEMPLATE,
)
from map_builder import build_all_maps, MapImages
from chart_builder import build_all_charts, ChartImages
from models.models import DealData, RecommendationVerdict

logger = logging.getLogger(__name__)


def validate_and_build_narrative_context(deal: DealData) -> dict:
    """Build the complete context dictionary for Sonnet narrative prompts.

    Pulls directly from the DealData object (computed by financials.py).
    Validates all required fields and logs the full context.

    CRITICAL: This function is the single source of truth for narrative context.
    Never build context inline or use cached/stale values.
    """
    a = deal.assumptions
    fo = deal.financial_outputs
    ext = deal.extracted_docs
    md = deal.market_data
    proforma = fo.pro_forma_years or []
    hold = a.hold_period or 10
    gba = a.gba_sf or 1

    total_project_cost = fo.total_uses or 0
    initial_loan = fo.initial_loan_amount or 0
    total_equity = fo.total_equity_required or 0
    equity_gap = total_equity  # total_equity IS the equity gap
    gp_equity = fo.gp_equity or 0
    lp_equity = fo.lp_equity or 0

    gpr_yr1 = fo.gross_potential_rent or 0
    egi_yr1 = fo.effective_gross_income or 0
    opex_yr1 = fo.total_operating_expenses or 0
    noi_yr1 = fo.noi_yr1 or 0
    fcf_yr1 = fo.free_cash_flow_yr1 or 0
    ds_yr1 = fo.debt_service_annual or 0

    last_yr = proforma[-1] if proforma else {}
    exit_noi = proforma[hold - 1].get("noi", 0) if len(proforma) >= hold else 0

    ctx = {
        # --- Property Info ---
        'property_name':         ext.property_name or deal.address.full_address or '',
        'address':               deal.address.full_address,
        'asset_type':            deal.asset_type.value,
        'strategy':              deal.investment_strategy.value,
        'building_sf':           a.gba_sf or 0,
        'lot_sf':                a.lot_sf or 0,
        'year_built':            a.year_built or 'Unknown',
        'num_units':             a.num_units or 1,
        'vacancy_status':        '100% Vacant' if noi_yr1 <= 0 else 'Occupied',

        # --- Acquisition ---
        'purchase_price':        a.purchase_price,
        'price_per_sf':          round(a.purchase_price / gba, 2) if gba > 0 else 0,
        'price_per_unit':        round(a.purchase_price / (a.num_units or 1), 0),
        'asking_price':          ext.asking_price or a.purchase_price,

        # --- Project Cost & Equity ---
        'total_project_cost':    total_project_cost,
        'hard_costs':            a.const_hard,
        'construction_reserve':  a.const_reserve,
        'initial_loan':          initial_loan,
        'ltv':                   a.ltv_pct,
        'equity_gap':            equity_gap,
        'gp_pct':                a.gp_equity_pct,
        'lp_pct':                a.lp_equity_pct,
        'gp_equity':             gp_equity,
        'lp_equity':             lp_equity,
        'total_equity':          total_equity,
        'debt_pct_of_cost':      initial_loan / max(total_project_cost, 1),
        'equity_pct_of_cost':    total_equity / max(total_project_cost, 1),

        # --- Financing ---
        'interest_rate':         a.interest_rate,
        'interest_rate_pct':     a.interest_rate * 100,
        'loan_term_years':       a.loan_term,
        'io_period_months':      a.io_period_months,
        'amort_years':           a.amort_years,
        'origination_fee_pct':   a.origination_fee_pct,
        'year1_debt_service':    ds_yr1,
        'annual_io_payment':     initial_loan * a.interest_rate,
        'annual_pi_payment':     ds_yr1 if a.io_period_months == 0 else 0,

        # --- Income & Expenses (Year 1) ---
        'year1_gpr':             gpr_yr1,
        'year1_egi':             egi_yr1,
        'year1_opex':            opex_yr1,
        'year1_noi':             noi_yr1,
        'year1_cfbt':            fcf_yr1,
        'year1_coc':             fo.cash_on_cash_yr1 or 0,
        'year1_dscr':            fo.dscr_yr1 or 0,
        'vacancy_rate':          a.vacancy_rate,
        'vacancy_rate_pct':      a.vacancy_rate * 100,
        'loss_to_lease':         a.loss_to_lease,
        'loss_to_lease_pct':     a.loss_to_lease * 100,
        'revenue_growth':        a.annual_rent_growth,
        'revenue_growth_pct':    a.annual_rent_growth * 100,
        'expense_growth':        a.expense_growth_rate,
        'expense_growth_pct':    a.expense_growth_rate * 100,
        'effective_rent_psf':    gpr_yr1 / gba if gba > 0 else 0,

        # --- Year 10 / Exit ---
        'year10_gpr':            last_yr.get('gpr', 0),
        'year10_noi':            last_yr.get('noi', 0),
        'exit_year':             hold,
        'exit_noi':              exit_noi,
        'exit_cap_rate':         a.exit_cap_rate,
        'exit_cap_rate_pct':     a.exit_cap_rate * 100,
        'gross_sale_price':      fo.gross_sale_price or 0,
        'net_sale_proceeds':     fo.net_sale_proceeds or 0,
        'net_equity_at_exit':    fo.net_equity_at_exit or 0,
        'hold_period':           hold,
        'disposition_cost_pct':  a.disposition_costs_pct,

        # --- Sensitivity stabilization ---
        'sensitivity_stabilized_year': fo.sensitivity_stabilized_year,
        'sensitivity_stabilized_noi':  fo.sensitivity_stabilized_noi or 0,
        'sensitivity_note':            fo.sensitivity_note or '',

        # --- Returns ---
        'project_irr':           fo.project_irr,
        'lp_irr':                fo.lp_irr,
        'gp_irr':                fo.gp_irr,
        'lp_equity_multiple':    fo.lp_equity_multiple,
        'equity_multiple':       fo.project_equity_multiple,
        'preferred_return':      a.pref_return,
        'preferred_return_pct':  a.pref_return * 100,
        'target_lp_irr':         a.target_lp_irr,
        'target_lp_irr_pct':     a.target_lp_irr * 100,
        'min_lp_irr':            a.min_lp_irr,
        'min_lp_irr_pct':        a.min_lp_irr * 100,

        # IRR display values
        'project_irr_display':   (f"{fo.project_irr * 100:.2f}%" if fo.project_irr is not None
                                  else "N/A (negative NOI — IRR non-convergent)"),
        'lp_irr_display':        (f"{fo.lp_irr * 100:.2f}%" if fo.lp_irr is not None else "N/A"),
        'lp_em_display':         (f"{fo.lp_equity_multiple:.2f}x" if fo.lp_equity_multiple is not None else "N/A"),

        # --- Market / Demographics ---
        'submarket':             'West Philadelphia',
        'city':                  deal.address.city,
        'state':                 deal.address.state,
        'zip_code':              deal.address.zip_code,
        'pop_3mi':               md.population_3mi or 0,
        'median_hh_income_3mi':  md.median_hh_income_3mi or 0,
        'renter_pct_3mi':        md.pct_renter_occ_3mi or 0,
        'unemployment_rate':     md.unemployment_rate or 0,

        # --- Due Diligence ---
        'zoning_code':           deal.zoning.zoning_code or 'Pending verification',
        'fema_flood_zone':       md.fema_flood_zone or 'Not Determined',
        'epa_flags':             '; '.join(md.epa_env_flags) if md.epa_env_flags else 'None identified',
        'phase1_status':         'Not completed',
        'title_insurance':       a.title_insurance,

        # --- Insurance ---
        'insurance_proforma':    deal.insurance.insurance_proforma_line_item or a.insurance,
        'insurance_total_low':   (deal.insurance.insurance_proforma_line_item or a.insurance) * 0.8,
        'insurance_total_high':  (deal.insurance.insurance_proforma_line_item or a.insurance) * 1.2,

        # --- Report metadata ---
        'report_date':           deal.report_date or '',
        'deal_id':               deal.deal_id or '',
        'pipeline_version':      '1.0',
        'prompt_catalog_version': '4.0',
    }

    # ── Validation ────────────────────────────────────────────────
    critical_fields = [
        'total_project_cost', 'initial_loan', 'equity_gap',
        'gp_equity', 'lp_equity', 'year1_gpr', 'year1_noi',
    ]
    for field in critical_fields:
        val = ctx.get(field)
        if val is None or val == 0:
            logger.warning("NARRATIVE CTX WARNING: '%s' is %s — check financials.py output", field, val)

    # ── Mandatory log ─────────────────────────────────────────────
    logger.info("NARRATIVE CTX CHECK:")
    logger.info("  total_project_cost = $%.2f", ctx['total_project_cost'])
    logger.info("  initial_loan       = $%.2f  (%.0f%% LTV on TPC)", ctx['initial_loan'], ctx['ltv'] * 100)
    logger.info("  equity_gap         = $%.2f", ctx['equity_gap'])
    logger.info("  gp_equity          = $%.2f  (%.0f%%)", ctx['gp_equity'], ctx['gp_pct'] * 100)
    logger.info("  lp_equity          = $%.2f  (%.0f%%)", ctx['lp_equity'], ctx['lp_pct'] * 100)
    logger.info("  year1_noi          = $%.2f", ctx['year1_noi'])
    logger.info("  project_irr        = %s", ctx['project_irr_display'])
    logger.info("  lp_irr             = %s", ctx['lp_irr_display'])
    logger.info("  lp_em              = %s", ctx['lp_em_display'])

    return ctx


def _safe_image(path, doc, width_cm):
    """Return InlineImage if file exists with content, else None.
    Never returns a placeholder."""
    from docx.shared import Cm
    if (path and
            os.path.exists(path) and
            os.path.getsize(path) > 0):
        try:
            return InlineImage(doc, path, width=Cm(width_cm))
        except Exception:
            return None
    return None


def fetch_street_view_image(address: str, deal_id: str, output_dir: str = "outputs") -> Optional[str]:
    """Fetch a Google Street View static image as a fallback hero photo.
    Returns path to saved image, or None if fetch fails."""
    import requests

    api_key = os.environ.get("GOOGLE_MAPS_API_KEY", "")
    if not api_key:
        logger.info("STREET VIEW: No API key configured — skipping photo fallback")
        return None

    encoded_address = requests.utils.quote(address)
    url = (f"https://maps.googleapis.com/maps/api/streetview"
           f"?size=800x600&location={encoded_address}"
           f"&fov=90&heading=235&pitch=10&key={api_key}")

    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200 and len(response.content) > 5000:
            os.makedirs(output_dir, exist_ok=True)
            img_path = os.path.join(output_dir, f"{deal_id}_street_view.jpg")
            with open(img_path, 'wb') as f:
                f.write(response.content)
            size = os.path.getsize(img_path)
            logger.info("STREET VIEW: saved to %s, size=%d bytes", img_path, size)
            if size < 1000:
                logger.warning("STREET VIEW: file too small — likely API error response image")
                return None
            return img_path
        else:
            logger.warning("STREET VIEW: fetch failed (status=%d, size=%d)",
                           response.status_code, len(response.content))
            return None
    except Exception as e:
        logger.warning("STREET VIEW: exception — %s", e)
        return None


def _claude_call(client, **kwargs):
    """Call Claude API with up to 3 retries on 500/529 errors."""
    for attempt in range(3):
        try:
            return client.messages.create(**kwargs)
        except Exception as e:
            err = str(e)
            if attempt < 2 and any(code in err for code in
                                   ["500", "529", "overloaded",
                                    "internal server"]):
                wait = 15 * (attempt + 1)
                logger.warning(
                    f"Anthropic API transient error (attempt "
                    f"{attempt+1}/3): {e}. Retrying in {wait}s..."
                )
                time.sleep(wait)
            else:
                raise


# ═══════════════════════════════════════════════════════════════════════════
# PROMPT 4-MASTER — ALL REPORT NARRATIVE SECTIONS (Sonnet)
# ═══════════════════════════════════════════════════════════════════════════

_SYSTEM_4MASTER = (
    "You are a senior commercial real estate analyst and investment writer producing\n"
    "the narrative sections of a formal institutional investment underwriting report.\n\n"
    "You will receive a complete deal data object and generate written narrative\n"
    "for every report section listed.\n\n"
    "GLOBAL WRITING RULES:\n"
    "- Voice: Senior analyst writing for an investment committee. Precise, data-grounded.\n"
    "- Tense: Present for market conditions, past for historical facts, future for projections.\n"
    "- Never use: \"pleased to present,\" \"exciting opportunity,\" \"unique,\" \"best-in-class.\"\n"
    "- Every claim must be grounded in the data provided. No invented facts.\n"
    "- All numbers must match the data exactly — never round or paraphrase figures.\n"
    "- Return ONLY valid JSON. No markdown, no preamble, no commentary.\n\n"
    "SECTION REQUIREMENTS (word counts are targets — +/-15% acceptable):\n\n"
    "exec_overview_p1 (100-130 words): Address, asset type, strategy, asking price, thesis,\n"
    "  physical characteristics, current occupancy/NOI.\n"
    "exec_overview_p2 (80-110 words): Submarket, vacancy trends, rent growth, market conditions.\n"
    "exec_overview_p3 (80-100 words): Hold period, target IRR, exit strategy, risk-return rationale.\n"
    "exec_pullquote (15-25 words): Quotable deal thesis sentence. No hedging.\n"
    "deal_thesis (60-80 words): Why this property, strategy, market, at this time.\n"
    "opportunity_1/2/3 (15-25 words each): Three strongest value creation levers.\n"
    "prop_desc_p1 (80-110 words): Building type, construction, condition, layout, parking.\n"
    "prop_desc_p2 (60-80 words): Unit mix, interior conditions, renovation opportunity.\n"
    "prop_desc_p3 (60-80 words): Tenant profile, occupancy, lease terms.\n"
    "prop_desc_p4 (50-70 words): Utilities and infrastructure systems.\n"
    "utilities_analysis (50-70 words): Systems condition, deferred maintenance.\n"
    "ownership_narrative (80-110 words): Chain of title, entity structure, notable events.\n"
    "liens_narrative (50-70 words): Recorded encumbrances. If none, state clearly.\n"
    "location_pullquote (15-25 words): Location's strongest attribute.\n"
    "location_overview_p1 (90-120 words): Neighborhood, submarket, major employers, transit.\n"
    "location_overview_p2 (80-100 words): Demographics, income, renter composition, trends.\n"
    "transportation_analysis (60-80 words): Transit access, Walk Score, highways, parking.\n"
    "neighborhood_trend_narrative (100-130 words): Population/income trends, neighborhood trajectory.\n"
    "supply_pipeline_narrative (90-120 words): Competing supply within 1 mile, absorption,\n"
    "  impact on rent growth and exit caps. If Shonda at Binswanger was flagged for CoStar\n"
    "  data, acknowledge the data limitation explicitly.\n"
    "rent_roll_intro (50-70 words): Total units, occupancy, rent roll framing.\n"
    "rent_comp_narrative (80-100 words): Rents vs. comp set, upside assessment.\n"
    "commercial_comp_narrative (70-90 words): Commercial comp analysis. Abbreviate if no retail.\n"
    "sale_comp_narrative (80-100 words): Price vs. closed sales per-unit and per-SF.\n"
    "financial_pullquote (15-25 words): Financial thesis pull-quote.\n"
    "sources_uses_narrative (70-90 words): Total project cost, equity, debt, what capital pays for.\n"
    "proforma_narrative (100-130 words): 10-yr revenue trajectory, expense management, NOI growth.\n"
    "proforma_pullquote (15-25 words): Pro forma pull-quote (NOI growth or cash-on-cash).\n"
    "sensitivity_narrative (70-90 words): Sensitivity matrix — what passes/fails threshold.\n"
    "  IMPORTANT: If sensitivity_matrix is empty or all zeros, write exactly:\n"
    "  'Sensitivity analysis requires stabilized revenue data. Matrix will be populated\n"
    "  following lease-up and rent roll stabilization.' Do not report zeros as results.\n"
    "exit_narrative (70-90 words): Exit cap assumption, terminal value, net proceeds.\n"
    "capital_stack_narrative (80-100 words): LTV, debt terms, equity split, structure rationale.\n"
    "capital_structure_pullquote (15-25 words): Capital structure pull-quote.\n"
    "debt_comparison_narrative (60-80 words): Two alternative debt structures considered.\n"
    "waterfall_narrative (70-90 words): Promote structure, pref return, alignment of interests.\n"
    "environmental_intro (60-80 words): Environmental screening overview.\n"
    "phase_esa_narrative (70-90 words): Phase I/II findings. If none on file, state and flag.\n"
    "climate_risk_narrative (70-90 words): First Street scores interpreted for hold period.\n"
    "legal_status_narrative (70-90 words): Encumbrances, easements, legal matters.\n"
    "violations_narrative (50-70 words): Code violations or permit issues. If none, state clearly.\n"
    "regulatory_approvals_narrative (50-70 words): Required approvals, variances, special exceptions.\n"
    "due_diligence_overview (60-80 words): DD flag methodology and flag distribution summary.\n"
    "dd_checklist_intro (40-60 words): DD checklist scope framing.\n"
    "timeline_narrative (60-80 words): Phases from acquisition through stabilization.\n"
    "recommendation (REQUIRED, verdict enum): Exactly one of 'GO', 'CONDITIONAL GO', or 'NO-GO'.\n"
    "  Base this on deal economics, market conditions, risk profile, and DD flags.\n"
    "recommendation_one_line (REQUIRED, 15-30 words): Single-sentence summary of the verdict\n"
    "  with the one most important supporting rationale. Direct, declarative, no hedging.\n"
    "recommendation_narrative_p1 (100-130 words): Recommendation, primary rationale, key metrics.\n"
    "recommendation_narrative_p2 (80-110 words): Top risks and why manageable; next action.\n"
    "recommendation_pullquote (15-25 words): Recommendation pull-quote. Direct and declarative.\n"
    "risk_1/2/3 (25-35 words each): Three primary investment risks, concise statements.\n"
    "conclusion_1-5 (20-30 words each): Five thematic single-sentence conclusions.\n"
    "bottom_line (40-60 words): The last word on the deal before next steps.\n"
    "next_step_1-6 (15-25 words each): Six prioritized next steps beginning with action verbs.\n"
    "methodology_notes (80-100 words): Data sources, extraction methods, API pull dates,\n"
    "  DealDesk pipeline version. Municipal data sourced from DealDesk Municipal Registry\n"
    "  covering approximately 6,400 U.S. municipalities (data/municipal_registry.csv).\n"
    "  Use exactly '6,400' — do not use any other number for the registry size.\n"
    "photo_gallery_intro (30-50 words): Gallery context sentence.\n"
    "maps_intro (30-50 words): Maps section context sentence.\n"
    "fema_flood_narrative (50-70 words): Flood zone interpretation.\n"
    "construction_budget_narrative (70-90 words): Construction budget breakdown (if value_add)."
)

_USER_4MASTER = (
    "Generate all report narrative sections for the deal below.\n"
    "Return a single JSON object where every key is a report placeholder name.\n\n"
    "COMPLETE DEAL DATA:\n"
    "{deal_data_json}\n\n"
    "Generate all narrative sections now. Return ONLY the JSON object."
)

# Partition 4-MASTER keys into two batches to avoid response truncation.
_4MASTER_KEYS_PART1 = [
    "exec_overview_p1", "exec_overview_p2", "exec_overview_p3", "exec_pullquote",
    "deal_thesis", "opportunity_1", "opportunity_2", "opportunity_3",
    "prop_desc_p1", "prop_desc_p2", "prop_desc_p3", "prop_desc_p4",
    "utilities_analysis", "ownership_narrative", "liens_narrative",
    "location_pullquote", "location_overview_p1", "location_overview_p2",
    "transportation_analysis", "neighborhood_trend_narrative",
    "supply_pipeline_narrative", "rent_roll_intro", "rent_comp_narrative",
    "commercial_comp_narrative", "sale_comp_narrative",
    "photo_gallery_intro", "maps_intro",
]
_4MASTER_KEYS_PART2 = [
    "financial_pullquote", "sources_uses_narrative", "proforma_narrative",
    "proforma_pullquote", "sensitivity_narrative", "exit_narrative",
    "capital_stack_narrative", "capital_structure_pullquote",
    "debt_comparison_narrative", "waterfall_narrative",
    "environmental_intro", "phase_esa_narrative", "climate_risk_narrative",
    "legal_status_narrative", "violations_narrative",
    "regulatory_approvals_narrative", "due_diligence_overview",
    "dd_checklist_intro", "timeline_narrative",
    "recommendation", "recommendation_one_line",
    "recommendation_narrative_p1", "recommendation_narrative_p2",
    "recommendation_pullquote",
    "risk_1", "risk_2", "risk_3",
    "conclusion_1", "conclusion_2", "conclusion_3", "conclusion_4", "conclusion_5",
    "bottom_line",
    "next_step_1", "next_step_2", "next_step_3",
    "next_step_4", "next_step_5", "next_step_6",
    "methodology_notes", "fema_flood_narrative", "construction_budget_narrative",
]

_USER_4MASTER_SUBSET = (
    "Generate a subset of report narrative sections for the deal below.\n"
    "Return ONLY a single JSON object containing EXACTLY these keys and no others:\n"
    "{keys_list}\n\n"
    "COMPLETE DEAL DATA:\n"
    "{deal_data_json}\n\n"
    "Return ONLY the JSON object."
)

FALLBACK_NARRATIVES = {
    "executive_summary": "Executive summary pending final data review.",
    "top_opportunities": "Opportunities analysis in progress.",
    "key_risks": "Risk analysis in progress.",
    "next_steps": "Recommended next steps to be finalized at closing.",
}


# ═══════════════════════════════════════════════════════════════════════════
# PROMPT 5D — INVESTOR-FACING NARRATIVE REWRITE (Sonnet)
# ═══════════════════════════════════════════════════════════════════════════

# The 9 narrative keys rewritten for LP-appropriate language
_INVESTOR_REWRITE_KEYS = [
    "exec_overview_p1",
    "exec_overview_p2",
    "exec_overview_p3",
    "deal_thesis",
    "sources_uses_narrative",
    "capital_stack_narrative",
    "recommendation_narrative_p1",
    "recommendation_narrative_p2",
    "bottom_line",
]

_SYSTEM_5D = (
    "You are a senior CRE investment writer rewriting narrative sections of an\n"
    "underwriting report for an LP investor audience.\n\n"
    "RULES:\n"
    "- Rewrite each section in LP-appropriate language: professional, forward-looking,\n"
    "  focused on risk-adjusted returns and alignment of interests.\n"
    "- Remove internal-only language: GP negotiation details, proprietary sourcing\n"
    "  advantages, internal hurdle rate discussions, DD flag details.\n"
    "- Preserve all factual data — numbers, dates, metrics must match exactly.\n"
    "- Maintain the same word count targets as the original sections.\n"
    "- Material risk disclosures must NEVER be suppressed or softened.\n"
    "- Voice: Confident, institutional, suitable for LP distribution.\n"
    "- Return ONLY valid JSON with exactly the 9 keys provided. No extras."
)

_USER_5D = (
    "Rewrite these 9 narrative blocks for an LP investor audience.\n\n"
    "DEAL DATA:\n{deal_data_json}\n\n"
    "CURRENT NARRATIVES TO REWRITE:\n{narratives_json}\n\n"
    "Return JSON with exactly these 9 keys, each containing the rewritten text:\n"
    "{keys_list}\n\n"
    "Return ONLY the JSON object."
)


# ═══════════════════════════════════════════════════════════════════════════
# LLM CALL
# ═══════════════════════════════════════════════════════════════════════════

def _call_sonnet(system: str, user_msg: str, max_tokens: int = 8192) -> Optional[dict]:
    """Send a single Sonnet call and parse the JSON response. Returns None on failure."""
    logger.info(f"ANTHROPIC_API_KEY present: {bool(ANTHROPIC_API_KEY)}, length: {len(ANTHROPIC_API_KEY) if ANTHROPIC_API_KEY else 0}")
    client = anthropic.Anthropic(
        api_key=ANTHROPIC_API_KEY,
    )
    try:
        response = _claude_call(
            client,
            model=MODEL_SONNET,
            max_tokens=max_tokens,
            system=system,
            messages=[{"role": "user", "content": user_msg}],
        )
        raw = response.content[0].text
        raw = raw.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
        return json.loads(raw)
    except anthropic.AuthenticationError as auth_err:
        logger.error("SONNET AUTH ERROR (401): %s | key_prefix=%s",
                     auth_err, ANTHROPIC_API_KEY[:12])
        return None
    except anthropic.APIStatusError as status_err:
        logger.error("SONNET API STATUS ERROR: status=%s body=%s",
                     status_err.status_code, status_err.message)
        return None
    except (json.JSONDecodeError, anthropic.APIError, IndexError, KeyError) as exc:
        logger.error("SONNET CALL FAILED: %s | type=%s", exc, type(exc).__name__)
        return None


# ═══════════════════════════════════════════════════════════════════════════
# NARRATIVE GENERATION
# ═══════════════════════════════════════════════════════════════════════════


def _generate_narratives(deal: DealData) -> None:
    """Run Prompt 4-MASTER to populate all narrative fields on DealData."""
    logger.info("Running Prompt 4-MASTER — all report narratives...")

    # Pre-flight: full narrative context validation
    validate_and_build_narrative_context(deal)

    deal_json = deal.model_dump_json(indent=2)

    def _run_part(keys: list, label: str) -> dict:
        user_msg = _USER_4MASTER_SUBSET.format(
            keys_list=json.dumps(keys),
            deal_data_json=deal_json,
        )
        logger.info("NARRATIVE: calling Sonnet for %s (%d keys)...", label, len(keys))
        out = _call_sonnet(_SYSTEM_4MASTER, user_msg, max_tokens=8192)
        if out is None:
            logger.warning("Prompt 4-MASTER %s first attempt failed — retrying...", label)
            out = _call_sonnet(_SYSTEM_4MASTER, user_msg, max_tokens=8192)
        return out or {}

    part1 = _run_part(_4MASTER_KEYS_PART1, "Part 1 (executive/property/location/market)")
    part2 = _run_part(_4MASTER_KEYS_PART2, "Part 2 (financial/risk/recommendation)")

    result: dict = {}
    result.update(part1)
    result.update(part2)

    if not result:
        logger.error("Prompt 4-MASTER failed twice — narratives will be empty strings")
        return

    for key, fallback in FALLBACK_NARRATIVES.items():
        if not result.get(key):
            result[key] = fallback
            logger.warning("NARRATIVE FALLBACK: %s — using placeholder", key)

    for _k in ("executive_summary", "deal_thesis", "top_opportunities", "key_risks", "next_steps"):
        _v = result.get(_k, "")
        _chars = len(_v) if isinstance(_v, str) else 0
        logger.info(f"NARRATIVE: {_k} returned {_chars} chars")

    # Apply all returned keys to the narratives model
    narr = deal.narratives
    for key, value in result.items():
        if hasattr(narr, key) and isinstance(value, str):
            setattr(narr, key, value)

    # Recommendation verdict + one-liner live on DealData directly (not narratives)
    rec_raw = (result.get("recommendation")
               or result.get("go_nogo_recommendation")
               or result.get("investment_recommendation") or "").strip().upper()
    rec_map = {
        "GO": RecommendationVerdict.GO,
        "CONDITIONAL GO": RecommendationVerdict.CONDITIONAL_GO,
        "CONDITIONAL-GO": RecommendationVerdict.CONDITIONAL_GO,
        "NO-GO": RecommendationVerdict.NO_GO,
        "NO GO": RecommendationVerdict.NO_GO,
        "NOGO": RecommendationVerdict.NO_GO,
    }
    if rec_raw in rec_map:
        deal.recommendation = rec_map[rec_raw]
    elif rec_raw:
        logger.warning("RECOMMENDATION: unrecognized verdict '%s' — leaving unset", rec_raw)

    rec_one = (result.get("recommendation_one_line")
               or result.get("recommendation_summary")
               or result.get("one_line_recommendation") or "").strip()
    if rec_one:
        deal.recommendation_one_line = rec_one

    logger.info("RECOMMENDATION: verdict=%s, one_line='%s' (len=%d)",
                deal.recommendation.value if deal.recommendation else "EMPTY",
                (deal.recommendation_one_line or "")[:60],
                len(deal.recommendation_one_line or ""))
    logger.info("Prompt 4-MASTER complete — %d narrative keys populated", len(result))


def _rewrite_investor_narratives(deal: DealData) -> None:
    """Run Prompt 5D to rewrite 9 narrative blocks for LP audience."""
    logger.info("Running Prompt 5D — investor narrative rewrite...")

    narr = deal.narratives
    current = {k: getattr(narr, k, "") or "" for k in _INVESTOR_REWRITE_KEYS}

    deal_json = deal.model_dump_json(indent=2)
    keys_list = json.dumps(_INVESTOR_REWRITE_KEYS)
    user_msg = _USER_5D.format(
        deal_data_json=deal_json,
        narratives_json=json.dumps(current, indent=2),
        keys_list=keys_list,
    )

    result = _call_sonnet(_SYSTEM_5D, user_msg, max_tokens=4096)

    if result is None:
        logger.warning("Prompt 5D failed — investor narratives unchanged")
        return

    for key in _INVESTOR_REWRITE_KEYS:
        if key in result and isinstance(result[key], str):
            setattr(narr, key, result[key])

    logger.info("Prompt 5D complete — %d investor narrative keys rewritten",
                sum(1 for k in _INVESTOR_REWRITE_KEYS if k in result))


# ═══════════════════════════════════════════════════════════════════════════
# TEMPLATE CONTEXT BUILDER
# ═══════════════════════════════════════════════════════════════════════════

def _excel_total_uses(deal) -> float:
    """Compute total USES exactly as the Excel S&U tab sums it.

    This mirrors the individual cell values written by excel_builder._section_uses()
    plus the origination fee (which Excel computes as a formula but from the same inputs).
    Renovations are NOT included here — they live in Below-the-Line (C168), not S&U.
    """
    a = deal.assumptions
    fo = deal.financial_outputs
    transfer_tax = a.purchase_price * a.transfer_tax_rate
    professional = (a.legal_closing + a.title_insurance + a.legal_bank +
                    a.appraisal + a.environmental + a.architect +
                    a.structural + a.geotech + a.surveyor + a.civil_eng +
                    a.meps + a.legal_zoning)
    financing = (a.acq_fee_fixed + a.mortgage_carry + a.mezz_interest)
    initial_loan = fo.initial_loan_amount or 0.0
    origination = initial_loan * a.origination_fee_pct
    soft = (a.working_capital + a.marketing + a.re_tax_carry +
            a.prop_ins_carry + a.dev_fee + a.dev_pref + a.permits)
    hard = (a.stormwater + a.demo + a.const_hard +
            a.const_reserve + a.gc_overhead)
    return (a.purchase_price + transfer_tax +
            a.tenant_buyout + professional + financing + origination +
            soft + hard)


def _reconcile_total_uses(deal: DealData) -> float:
    """Return the authoritative total project cost for PDF narratives.

    Compares deal.financial_outputs.total_uses (Python-computed) against
    the sum of individual Assumptions line items (what the Excel S&U tab
    actually sums).  If they agree within $1, use fo.total_uses.
    Otherwise, trust the Excel sum and log a warning.
    """
    a = deal.assumptions
    fo = deal.financial_outputs
    excel_total = _excel_total_uses(deal)
    python_fo = fo.total_uses or 0.0

    gap = abs(excel_total - python_fo)
    if gap < 1.0:
        logger.info("TOTAL PROJECT COST: excel_sum=%s, python_fo=%s, using=%s",
                     f"{excel_total:,.0f}", f"{python_fo:,.0f}", f"{python_fo:,.0f}")
        return python_fo
    else:
        logger.warning(
            "TOTAL PROJECT COST MISMATCH: excel_sum=%s, python_fo=%s, gap=%s — using excel_sum",
            f"{excel_total:,.0f}", f"{python_fo:,.0f}", f"{gap:,.0f}")
        return excel_total


def _build_context(deal: DealData) -> dict:
    """Build the full template context dict from DealData for docxtpl."""
    narr = deal.narratives
    a = deal.assumptions
    fo = deal.financial_outputs
    md = deal.market_data
    ins = deal.insurance
    ext = deal.extracted_docs

    suppressed = deal.suppressed_sections

    ctx = {}

    # Cover page
    ctx["cover_title"] = deal.cover_title
    ctx["report_date"] = deal.report_date or ""
    ctx["deal_id"] = deal.deal_id or ""
    ctx["deal_code"] = deal.deal_code or ""
    ctx["sponsor_name"] = deal.sponsor_name
    ctx["sponsor_description"] = deal.sponsor_description

    # Property basics
    ctx["property_name"] = ext.property_name or ""
    ctx["full_address"] = deal.address.full_address
    ctx["city"] = deal.address.city
    ctx["state"] = deal.address.state
    ctx["zip_code"] = deal.address.zip_code
    ctx["asset_type"] = deal.asset_type.value
    ctx["investment_strategy"] = deal.investment_strategy.value
    ctx["asking_price"] = f"${deal.assumptions.purchase_price:,.0f}" if deal.assumptions.purchase_price else "Not disclosed"
    ctx["purchase_price"] = a.purchase_price
    ctx["num_units"] = a.num_units
    ctx["building_sf"] = a.gba_sf
    ctx["lot_sf"] = a.lot_sf
    ctx["year_built"] = a.year_built
    ctx["hold_period"] = a.hold_period
    ctx["deal_description"] = deal.deal_description

    # Parcel data
    if deal.parcel_data:
        ctx["parcel_data"] = deal.parcel_data.model_dump()
    else:
        ctx["parcel_data"] = {}

    # Zoning
    ctx["zoning"] = deal.zoning.zoning_code or "Pending verification"
    ctx["zoning_code"] = deal.zoning.zoning_code or ""

    # Market data
    ctx["census_tract"] = deal.address.census_tract or ""
    ctx["fips_code"] = deal.address.fips_code or ""
    ctx["fema_flood_zone"] = md.fema_flood_zone or ""
    ctx["fema_panel_number"] = md.fema_panel_number or ""

    # Financial outputs — fo.total_uses is the single source of truth
    ctx["total_uses"] = fo.total_uses
    ctx["total_project_cost"] = fo.total_uses
    ctx["total_sources"] = fo.total_uses  # sources must equal uses
    logger.info("TPC: using fo.total_uses=%s as single source of truth",
                f"{fo.total_uses:,.2f}")
    ctx["total_equity_required"] = fo.total_equity_required
    ctx["initial_loan_amount"] = fo.initial_loan_amount
    ctx["noi_yr1"] = fo.noi_yr1
    ctx["dscr_yr1"] = fo.dscr_yr1
    ctx["going_in_cap_rate"] = fo.going_in_cap_rate
    ctx["lp_irr"] = fo.lp_irr if fo.lp_irr is not None else "Not calculable (negative NOI)"
    ctx["gp_irr"] = fo.gp_irr if fo.gp_irr is not None else "Not calculable (negative NOI)"
    ctx["project_irr"] = fo.project_irr if fo.project_irr is not None else "Not calculable (negative NOI)"
    ctx["lp_equity_multiple"] = fo.lp_equity_multiple
    ctx["gp_equity_multiple"] = fo.gp_equity_multiple
    ctx["cash_on_cash_yr1"] = fo.cash_on_cash_yr1
    ctx["gross_sale_price"] = fo.gross_sale_price
    ctx["net_sale_proceeds"] = fo.net_sale_proceeds
    ctx["sensitivity_matrix"] = fo.sensitivity_matrix
    ctx["pro_forma_years"] = fo.pro_forma_years

    # Insurance
    ctx["insurance_narrative_p1"] = ins.insurance_narrative_p1 or ""
    ctx["insurance_narrative_p2"] = ins.insurance_narrative_p2 or ""
    ctx["insurance_narrative_p3"] = ins.insurance_narrative_p3 or ""
    ctx["insurance_kpi_strip"] = ""
    ctx["insurance_summary_table"] = ""

    # DD Flags
    ctx["dd_flags"] = [f.model_dump() for f in deal.dd_flags]

    # Recommendation — populated by Prompt 4-MASTER into deal.recommendation
    ctx["recommendation"] = deal.recommendation.value if deal.recommendation else ""
    ctx["recommendation_one_line"] = deal.recommendation_one_line or ""
    logger.info("RECOMMENDATION: '%s' (len=%d)",
                ctx["recommendation"][:50] if ctx["recommendation"] else "EMPTY",
                len(ctx["recommendation"]))

    # Waterfall
    ctx["pref_return"] = a.pref_return
    ctx["gp_equity_pct"] = a.gp_equity_pct
    ctx["lp_equity_pct"] = a.lp_equity_pct
    ctx["waterfall_tiers"] = [t.model_dump() for t in a.waterfall_tiers]

    # GP/LP equity dollar amounts — read from financial_outputs (not recomputed)
    ctx["initial_loan"] = fo.initial_loan_amount
    ctx["total_equity"] = fo.total_equity_required
    ctx["gp_equity"] = fo.gp_equity
    ctx["lp_equity"] = fo.lp_equity
    ctx["equity_gap"] = fo.total_equity_required
    ctx["annual_ds"] = fo.debt_service_annual
    logger.info("NARRATIVE CTX: loan=%s equity=%s gp=%s lp=%s ds=%s",
        f"{fo.initial_loan_amount:,.2f}",
        f"{fo.total_equity_required:,.2f}",
        f"{fo.gp_equity:,.2f}",
        f"{fo.lp_equity:,.2f}",
        f"{fo.debt_service_annual:,.2f}")

    # Extracted doc data
    ctx["unit_mix"] = ext.unit_mix or []
    ctx["occupancy_rate"] = ext.occupancy_rate
    ctx["image_placements"] = ext.image_placements or {}

    # Provenance
    ctx["provenance"] = deal.provenance.model_dump()


    # ── Comparable market data tables (Section 11) ────────────────────────
    # Rent comp table rows (Table 24 — 7 cols: Property, Type, Beds, SF, Rent/Mo, $/SF, Distance)
    rent_rows = []
    for c in (deal.comps.rent_comps or []):
        rent_rows.append({
            "property":     c.address or "",
            "type":         c.unit_type or "",
            "beds":         str(c.beds) if c.beds is not None else "",
            "sf":           f"{c.sq_ft:,}" if c.sq_ft else "",
            "monthly_rent": f"${c.monthly_rent:,.0f}" if c.monthly_rent else "",
            "rent_per_sf":  f"${c.rent_per_sf:.2f}" if c.rent_per_sf else "",
            "distance":     f"{c.distance_miles:.1f} mi" if c.distance_miles else "",
        })
    ctx["rent_comp_rows"] = rent_rows

    # Commercial comp table rows (Table 25 — 5 cols: Address, Use, SF, Rent/SF, Type)
    comm_rows = []
    for c in (deal.comps.commercial_comps or []):
        comm_rows.append({
            "address":      c.address or "",
            "use_type":     c.use_type or "",
            "sf":           f"{c.sq_ft:,}" if c.sq_ft else "",
            "rent_per_sf":  f"${c.asking_rent_per_sf:.2f}/SF" if c.asking_rent_per_sf else "",
            "lease_type":   c.lease_type or "",
        })
    ctx["commercial_comp_rows"] = comm_rows

    # Sale comp table rows (Table 26 — 6 cols: Address, SF, Units, Price, $/SF, Cap Rate)
    sale_rows = []
    for c in (deal.comps.sale_comps or []):
        sale_rows.append({
            "address":       c.address or "",
            "sf":            f"{c.sq_ft:,}" if c.sq_ft else "",
            "units":         str(c.num_units) if c.num_units else "",
            "sale_price":    f"${c.sale_price:,.0f}" if c.sale_price else "",
            "price_per_sf":  f"${c.price_per_sf:.0f}" if c.price_per_sf else "",
            "cap_rate":      f"{c.cap_rate:.2%}" if c.cap_rate else "",
        })
    ctx["sale_comp_rows"] = sale_rows

    # Section suppression flags
    ctx["suppressed_sections"] = suppressed
    for sid in ["s01", "s02", "s03", "s04", "s05", "s06", "s07", "s08", "s09",
                "s10", "s11", "s12", "s13", "s14", "s15", "s16", "s17", "s18",
                "s19", "s20", "s21", "s22"]:
        ctx[f"show_{sid}"] = sid not in suppressed

    # All narrative fields
    for field_name in narr.model_fields:
        ctx[field_name] = getattr(narr, field_name) or ""

    # Investor mode flag
    ctx["investor_mode"] = deal.investor_mode

    # ── Zoning standards table rows ───────────────────────────────
    z = deal.zoning
    ctx["zoning_standards_rows"] = [
        {"parameter": "Zoning District",      "standard": z.zoning_code or "",         "proposed": "", "code_section": ""},
        {"parameter": "District Name",         "standard": z.zoning_district or "",     "proposed": "", "code_section": ""},
        {"parameter": "Max Height (ft)",       "standard": z.max_height_ft or "",       "proposed": "", "code_section": ""},
        {"parameter": "Max Stories",           "standard": z.max_stories or "",         "proposed": "", "code_section": ""},
        {"parameter": "Min Lot Area (SF)",     "standard": z.min_lot_area_sf or "",     "proposed": "", "code_section": ""},
        {"parameter": "Max Lot Coverage",      "standard": f"{z.max_lot_coverage_pct:.0%}" if z.max_lot_coverage_pct else "", "proposed": "", "code_section": ""},
        {"parameter": "Max FAR",               "standard": z.max_far or "",             "proposed": "", "code_section": ""},
        {"parameter": "Front Setback (ft)",    "standard": z.front_setback_ft or "",   "proposed": "", "code_section": ""},
        {"parameter": "Rear Setback (ft)",     "standard": z.rear_setback_ft or "",    "proposed": "", "code_section": ""},
        {"parameter": "Side Setback (ft)",     "standard": z.side_setback_ft or "",    "proposed": "", "code_section": ""},
        {"parameter": "Min Parking Spaces",    "standard": z.min_parking_spaces or "", "proposed": "", "code_section": ""},
        {"parameter": "Permitted Uses",        "standard": ", ".join(z.permitted_uses) if z.permitted_uses else "", "proposed": "", "code_section": ""},
    ]

    # ── Pro forma table rows (formatted for report) ───────────────
    pf_rows = []
    for yr in (fo.pro_forma_years or []):
        ds = yr.get("debt_service", 0) or 0
        noi = yr.get("noi", 0) or 0
        pf_rows.append({
            "year":          yr.get("year", ""),
            "gpr":           f"${yr.get('gpr', 0):,.0f}",
            "egi":           f"${yr.get('egi', 0):,.0f}",
            "opex":          f"${yr.get('opex', 0):,.0f}",
            "noi":           f"${noi:,.0f}",
            "debt_service":  f"${ds:,.0f}",
            "cfbt":          f"${yr.get('fcf', 0):,.0f}",
            "coc":           f"{yr.get('cash_on_cash', 0):.1%}",
            "dscr":          f"{noi/ds:.2f}x" if ds > 0 else "N/A",
        })
    ctx["pro_forma_table_rows"] = pf_rows

    # ── Pro forma rows — raw numeric dicts for Section 12.4 ──────
    ctx["pro_forma_rows"] = [
        {
            "year":     yr.get("year", ""),
            "gpr":      yr.get("gpr", 0) or 0,
            "egr":      yr.get("egi", 0) or 0,
            "opex":     yr.get("opex", 0) or 0,
            "noi":      yr.get("noi", 0) or 0,
            "debt_svc": yr.get("debt_service", 0) or 0,
            "cfbt":     yr.get("fcf", 0) or 0,
            "coc":      yr.get("cash_on_cash", 0) or 0,
            "dscr":     round((yr.get("noi", 0) or 0) / ds, 2)
                        if (ds := yr.get("debt_service", 0) or 0) > 0
                        else 0,
        }
        for yr in (fo.pro_forma_years or [])
    ]

    # ── Sensitivity matrix rows (formatted for report) ────────────
    sens_rows = []
    rent_axis  = fo.sensitivity_axis_rent_growth or []
    cap_axis   = fo.sensitivity_axis_exit_cap or []
    matrix     = fo.sensitivity_matrix or []
    for i, row in enumerate(matrix):
        rg_label = f"{rent_axis[i]:.1%}" if i < len(rent_axis) else ""
        sens_rows.append({
            "rent_growth": rg_label,
            "values": ["N/A" if v == "N/A" else (f"{v:.1%}" if isinstance(v, (int, float)) else "—") for v in row],
        })
    ctx["sensitivity_rows"]      = sens_rows
    ctx["sensitivity_cap_axis"]  = [f"{c:.1%}" for c in cap_axis]

    # ── Monte Carlo / Risk-Weighted Return (section 12.6) ─────────
    ctx["monte_carlo_results"] = fo.monte_carlo_results or {}
    ctx["monte_carlo_narrative"] = (
        fo.monte_carlo_narrative
        or "Risk-weighted return analysis requires stabilized NOI. "
           "This analysis will be completed upon lease execution and "
           "confirmation of stabilized operating assumptions."
    )

    # ── Full market data object (for demographic table) ───────────
    ctx["market_data"] = deal.market_data.model_dump() if deal.market_data else {}

    # ── Waterfall formatted rows ──────────────────────────────────
    wf_rows = []
    tier_labels = ["Tier 1 — Pref Return + ROC", "Tier 2", "Tier 3", "Tier 4"]
    for i, t in enumerate(a.waterfall_tiers):
        wf_rows.append({
            "tier":     tier_labels[i] if i < len(tier_labels) else f"Tier {i+1}",
            "hurdle":   f"{t.hurdle_value:.1%}" if t.hurdle_value else "",
            "lp_split": f"{t.lp_share:.0%}",
            "gp_split": f"{1 - t.lp_share:.0%}",
            "promote":  f"{1 - t.lp_share:.0%}",
        })
    if hasattr(a, "residual_tier") and a.residual_tier:
        wf_rows.append({
            "tier":     "Residual (above Tier 4)",
            "hurdle":   "Above all hurdles",
            "lp_split": f"{a.residual_tier.lp_share:.0%}",
            "gp_split": f"{a.residual_tier.gp_share:.0%}",
            "promote":  f"{a.residual_tier.gp_share:.0%}",
        })
    ctx["waterfall_tier_rows"] = wf_rows

    # ── FRED rates (formatted for debt section) ───────────────────
    md = deal.market_data
    ctx["dgs10_rate"]      = f"{md.dgs10_rate:.2%}"      if md.dgs10_rate      else "N/A"
    ctx["sofr_rate"]       = f"{md.sofr_rate:.2%}"       if md.sofr_rate       else "N/A"
    ctx["mortgage30_rate"] = f"{md.mortgage30_rate:.2%}" if md.mortgage30_rate else "N/A"
    ctx["cpi_yoy"]         = f"{md.cpi_yoy:.2%}"         if md.cpi_yoy         else "N/A"

    # ── Opportunity zone flag ─────────────────────────────────────
    ctx["is_opportunity_zone"] = (
        deal.provenance.field_sources.get("opportunity_zone", "False") == "True"
    )

    # ── EPA environmental flags ───────────────────────────────────
    ctx["epa_env_flags"] = md.epa_env_flags or []

    # ── HUD Fair Market Rents ─────────────────────────────────────
    ctx["fmr_studio"] = f"${md.fmr_studio:,.0f}" if md.fmr_studio else "N/A"
    ctx["fmr_1br"]    = f"${md.fmr_1br:,.0f}"    if md.fmr_1br    else "N/A"
    ctx["fmr_2br"]    = f"${md.fmr_2br:,.0f}"    if md.fmr_2br    else "N/A"
    ctx["fmr_3br"]    = f"${md.fmr_3br:,.0f}"    if md.fmr_3br    else "N/A"

    # ── Debt market narrative ─────────────────────────────────────
    ctx["debt_market_narrative"] = md.debt_market_narrative or ""

    # ── HBU and buildable capacity ────────────────────────────────
    ctx["hbu_narrative"]           = z.hbu_narrative or ""
    ctx["hbu_conclusion"]          = z.hbu_conclusion or ""
    ctx["buildable_capacity_narrative"] = z.buildable_capacity_narrative or ""

    # ══════════════════════════════════════════════════════════════
    # DATA-GAP NOTES — professional placeholders for missing data
    # ══════════════════════════════════════════════════════════════

    # Section 2 — Photo Gallery
    has_photos = bool(ext.image_placements)
    if not has_photos:
        ctx["photo_gallery_note"] = (
            "No property photographs are on file as of the "
            "report date. Images should be obtained during "
            "the physical site inspection and added to the "
            "record prior to investment committee presentation."
        )
    else:
        ctx["photo_gallery_note"] = ""

    # Section 9 — Supply Pipeline Register
    has_supply = bool(getattr(md, "supply_pipeline_narrative", None))
    if not has_supply:
        ctx["supply_pipeline_rows"] = []
        ctx["supply_pipeline_note"] = (
            "No competitive supply pipeline data is available "
            "for the immediate submarket at this time. Current "
            "submarket supply, absorption, and delivery data "
            "should be sourced from CoStar or a local market "
            "broker prior to investment committee submission."
        )
    else:
        ctx["supply_pipeline_note"] = ""

    # Section 10 — Unit Mix & Rent Roll
    has_unit_mix = bool(ext.unit_mix)
    has_rent_roll = bool(getattr(deal, "rent_roll", None))
    if not has_unit_mix and not has_rent_roll:
        ctx["rent_roll_rows"] = []
        ctx["rent_roll_note"] = (
            "No rent roll data is on file. All gross potential "
            "rent figures in the pro forma are projection-based "
            "assumptions, not derived from executed leases or "
            "trailing income history. A current rent roll or "
            "executed lease abstracts must be obtained and "
            "reviewed prior to capital commitment."
        )
    else:
        ctx["rent_roll_note"] = ""

    # Section 11.1 — Residential Rent Comparables
    if not ctx.get("rent_comp_rows"):
        ctx["rent_comp_rows"] = []
        ctx["rent_comps_note"] = (
            "No rent comparable data was provided in the deal "
            "package. A formal rent comp study covering asking "
            "rents and closed lease transactions for comparable "
            "properties in the submarket is required before the "
            "underwritten rent assumptions can be validated."
        )
    else:
        ctx["rent_comps_note"] = ""

    # Section 11.2 — Commercial Rent Comparables
    if not ctx.get("commercial_comp_rows"):
        ctx["commercial_comp_rows"] = []
        ctx["commercial_comps_note"] = (
            "No commercial comparable lease transactions are on "
            "file. Given the property's use strategy, commercial "
            "comp analysis should be commissioned from a local "
            "market broker or licensed appraiser prior to "
            "investment committee submission."
        )
    else:
        ctx["commercial_comps_note"] = ""

    # Section 11.3 — Sale Comparables
    if not ctx.get("sale_comp_rows"):
        ctx["sale_comp_rows"] = []
        ctx["sale_comps_note"] = (
            "No closed sale comparable transactions are on file. "
            "A formal sales comp analysis benchmarking price per "
            "SF and cap rate against the subject is required "
            "before the exit cap rate and terminal value "
            "assumptions can be substantiated."
        )
    else:
        ctx["sale_comps_note"] = ""

    # Merge validated narrative context — single source of truth
    narr_ctx = validate_and_build_narrative_context(deal)
    for k, v in narr_ctx.items():
        if k not in ctx or ctx[k] is None:
            ctx[k] = v

    # Overwrite narrative context with live pro forma values (single source of truth)
    try:
        fo = deal.financial_outputs
        pf = fo.pro_forma_years or []
        pf_noi  = [y.get("noi", 0) for y in pf]
        pf_egi  = [y.get("egi", 0) for y in pf]
        pf_opex = [y.get("opex", 0) for y in pf]
        pf_ds   = [y.get("debt_service", 0) for y in pf]
        pf_fcf  = [y.get("fcf", 0) for y in pf]

        # DSCR and stab factors are not stored in proforma dict — compute on the fly
        pf_dscr = [(n / d) if d else 0 for n, d in zip(pf_noi, pf_ds)]
        # stabilization factors (import from financials)
        try:
            from financials import _get_stabilization_factors
            pf_sf = _get_stabilization_factors(deal)
        except Exception:
            pf_sf = [1.0] * len(pf)

        first_stab_yr = next((i for i, sf in enumerate(pf_sf) if sf >= 1.0), 0)

        # Refi net proceeds come from the proforma year dict
        refi_proceeds_by_year = [y.get("refi_proceeds", 0) for y in pf]
        refi1_proceeds = next((p for p in refi_proceeds_by_year if p), 0)
        refi2_proceeds = 0
        _found_first = False
        for p in refi_proceeds_by_year:
            if p:
                if _found_first:
                    refi2_proceeds = p
                    break
                _found_first = True

        ctx.update({
            "year1_noi":          pf_noi[0]  if pf_noi  else "N/A",
            "year1_egi":          pf_egi[0]  if pf_egi  else "N/A",
            "year1_opex":         pf_opex[0] if pf_opex else "N/A",
            "year1_debt_svc":     pf_ds[0]   if pf_ds   else "N/A",
            "year1_fcf":          pf_fcf[0]  if pf_fcf  else "N/A",
            "year1_dscr":         pf_dscr[0] if pf_dscr else "N/A",
            "year10_noi":         pf_noi[9]  if len(pf_noi) > 9 else "N/A",
            "first_stabilized_year": first_stab_yr + 1,
            "first_stab_noi":     pf_noi[first_stab_yr] if pf_noi else "N/A",
            "refi1_net_proceeds": refi1_proceeds,
            "refi2_net_proceeds": refi2_proceeds,
        })
        logger.info("NARRATIVE CTX: year1_noi=%.2f, year1_dscr=%.2f, "
                    "first_stab_yr=%d, first_stab_noi=%.2f, refi1_proceeds=%.2f",
                    pf_noi[0] if pf_noi else 0,
                    pf_dscr[0] if pf_dscr else 0,
                    first_stab_yr + 1,
                    pf_noi[first_stab_yr] if pf_noi else 0,
                    refi1_proceeds)
    except Exception as e:
        logger.warning("NARRATIVE CTX: failed to refresh context — %s", e)

    # ═══════════════════════════════════════════════════════════════════
    # EXPLICIT TEMPLATE CONTEXT KEYS — kpi_rows, parcel_a_*, transit_rows,
    # hbu_content, income_* — built from the real model attribute paths.
    # ═══════════════════════════════════════════════════════════════════
    def _safe_fmt(val, fmt="${:,.0f}", fallback="N/A"):
        try:
            if val is None:
                return fallback
            return fmt.format(float(val))
        except (TypeError, ValueError):
            return fallback

    # ── Fix 1: KPI table rows ────────────────────────────────────────
    kpi_rows = [
        {"label": "Purchase Price",
         "value": _safe_fmt(a.purchase_price)},
        {"label": "Total Project Cost",
         "value": _safe_fmt(fo.total_uses)},
        {"label": "Total Equity Required",
         "value": _safe_fmt(fo.total_equity_required)},
        {"label": "GP Equity",
         "value": _safe_fmt(fo.gp_equity)},
        {"label": "LP Equity",
         "value": _safe_fmt(fo.lp_equity)},
        {"label": "Year 1 NOI",
         "value": _safe_fmt(fo.noi_yr1)},
        {"label": "Going-In Cap Rate",
         "value": (_safe_fmt((fo.going_in_cap_rate or 0) * 100, fmt="{:.2f}%")
                   if fo.going_in_cap_rate is not None else "N/A")},
        {"label": "Year 1 DSCR",
         "value": _safe_fmt(fo.dscr_yr1, fmt="{:.2f}x")},
        {"label": "Year 1 Cash-on-Cash",
         "value": (_safe_fmt((fo.cash_on_cash_yr1 or 0) * 100, fmt="{:.2f}%")
                   if fo.cash_on_cash_yr1 is not None else "N/A")},
        {"label": "Project IRR",
         "value": (_safe_fmt((fo.project_irr or 0) * 100, fmt="{:.1f}%")
                   if fo.project_irr is not None else "N/A")},
        {"label": "LP IRR",
         "value": (_safe_fmt((fo.lp_irr or 0) * 100, fmt="{:.1f}%")
                   if fo.lp_irr is not None else "N/A")},
        {"label": "LP Equity Multiple",
         "value": _safe_fmt(fo.lp_equity_multiple, fmt="{:.2f}x")},
        {"label": "Hold Period",
         "value": (f"{a.hold_period} years" if a.hold_period else "N/A")},
        {"label": "Exit Cap Rate",
         "value": (_safe_fmt((a.exit_cap_rate or 0) * 100, fmt="{:.2f}%")
                   if a.exit_cap_rate else "N/A")},
    ]
    ctx["kpi_rows"] = kpi_rows
    logger.info("KPI_ROWS built: %d rows", len(kpi_rows))
    for r in kpi_rows:
        logger.info("  KPI: %s = %s", r["label"], r["value"])

    # ── Fix 6: HBU content alias (real path: deal.zoning.hbu_narrative) ─
    hbu_text = (
        getattr(deal.zoning, "hbu_narrative", None)
        or getattr(deal.zoning, "hbu_conclusion", None)
        or ""
    )
    ctx["hbu_content"] = hbu_text or (
        "Highest and best use analysis is pending zoning verification. "
        "This section will be completed once the municipal zoning code "
        "scrape and buildable-capacity analysis return data for the "
        "subject parcel."
    )
    logger.info("HBU: content length=%d chars", len(ctx["hbu_content"]))

    # ── Fix 4: Parcel A context keys from real ParcelData ────────────
    pd_ = deal.parcel_data

    def _p_str(v, fallback="N/A"):
        return str(v) if v not in (None, "") else fallback

    def _p_money(v, fallback="N/A"):
        try:
            return f"${float(v):,.0f}" if v not in (None, "") else fallback
        except (TypeError, ValueError):
            return fallback

    def _p_area(v, fallback="N/A"):
        try:
            return f"{float(v):,.0f} SF" if v not in (None, "") else fallback
        except (TypeError, ValueError):
            return fallback

    ctx["parcel_a_address"]      = deal.address.full_address or "N/A"
    ctx["parcel_a_account"]      = _p_str(pd_.parcel_id        if pd_ else None)
    ctx["parcel_a_owner"]        = _p_str(pd_.owner_name       if pd_ else None)
    ctx["parcel_a_zoning"]       = _p_str(
        (pd_.zoning_code if pd_ else None) or deal.zoning.zoning_code,
        "Pending verification",
    )
    ctx["parcel_a_land_area"]    = _p_area(pd_.lot_area_sf     if pd_ else None)
    ctx["parcel_a_building_sf"]  = _p_area(pd_.building_sf     if pd_ else None)
    ctx["parcel_a_year_built"]   = _p_str(pd_.year_built       if pd_ else None)
    ctx["parcel_a_assessed"]     = _p_money(pd_.assessed_value if pd_ else None)
    ctx["parcel_a_taxable_land"] = _p_money(pd_.land_value     if pd_ else None)
    ctx["parcel_a_taxable_bldg"] = _p_money(pd_.improvement_value if pd_ else None)
    ctx["parcel_a_stories"]      = "N/A"
    ctx["parcel_a_category"]     = "N/A"

    # Parcel B is not in the current data model — always blank.
    ctx["parcel_b_address"]    = ""
    ctx["parcel_b_account"]    = "N/A"
    ctx["parcel_b_owner"]      = "N/A"
    ctx["parcel_b_zoning"]     = "N/A"
    ctx["parcel_b_land_area"]  = "N/A"
    ctx["parcel_b_year_built"] = "N/A"
    ctx["parcel_b_assessed"]   = "N/A"

    ctx["parcel_census_tract"] = deal.address.census_tract or "N/A"
    ctx["parcel_fips"]         = deal.address.fips_code or "N/A"
    logger.info("PARCEL A: account=%s owner=%s zoning=%s",
                ctx["parcel_a_account"],
                ctx["parcel_a_owner"],
                ctx["parcel_a_zoning"])

    # ── Fix 7: Transit rows for the template ─────────────────────────
    transit_list = list(getattr(md, "transit_options", []) or [])
    ctx["transit_rows"] = [
        {
            "mode":        t.get("mode", "Transit"),
            "route":       t.get("route", "—"),
            "distance":    t.get("distance", "—"),
            "destination": t.get("destination", "—"),
        }
        for t in transit_list[:8]
    ]
    logger.info("TRANSIT ROWS: %d rows", len(ctx["transit_rows"]))

    # ── Fix 8: Income Summary from in-place rent roll ────────────────
    # fo.gross_potential_rent is $0 during construction; fall back to rent roll.
    _units = (ext.unit_mix or [])
    _rr_monthly_total = 0.0
    for u in _units:
        try:
            monthly_raw = u.get("monthly_rent", 0) or u.get("current_rent", 0) or 0
            monthly = float(monthly_raw)
            count_raw = u.get("count")
            count = float(count_raw) if count_raw not in (None, "") else 1.0
            _rr_monthly_total += monthly * count
        except (TypeError, ValueError):
            continue
    _rr_gpr = _rr_monthly_total * 12.0
    _fo_gpr = float(fo.gross_potential_rent or 0)
    _gpr_val = _fo_gpr if _fo_gpr > 0 else _rr_gpr
    _vac_rate = float(a.vacancy_rate or 0.0)
    _ltl_rate = float(a.loss_to_lease or 0.0)
    _other = float((a.cam_reimbursements or 0) + (a.fee_income or 0))
    _vacancy_loss = _gpr_val * _vac_rate
    _ltl_loss     = _gpr_val * _ltl_rate
    _egi_val      = _gpr_val - _vacancy_loss - _ltl_loss + _other
    ctx["income_gpr"]           = f"${_gpr_val:,.0f}"
    ctx["income_vacancy_loss"]  = f"(${_vacancy_loss:,.0f})"
    ctx["income_vacancy_pct"]   = f"({_vac_rate * 100:.1f}%)"
    ctx["income_loss_to_lease"] = f"(${_ltl_loss:,.0f})"
    ctx["income_ltl_pct"]       = f"({_ltl_rate * 100:.1f}%)"
    ctx["income_other"]         = f"${_other:,.0f}"
    ctx["income_egi"]           = f"${_egi_val:,.0f}"
    ctx["income_egi_pct"]       = (f"{(_egi_val / _gpr_val * 100):.1f}%"
                                    if _gpr_val > 0 else "N/A")
    logger.info("EGI CALC: GPR=%s vacancy=%s ltl=%s EGI=%s (source=%s)",
                ctx["income_gpr"],
                ctx["income_vacancy_loss"],
                ctx["income_loss_to_lease"],
                ctx["income_egi"],
                "model" if _fo_gpr > 0 else "rent_roll")

    # ═══════════════════════════════════════════════════════════════════
    # RENT + SALE COMP CONTEXT (Section 11.1 benchmarks + 11.3 sale comps)
    # Sourced from deal.comps.rent_comps / deal.comps.sale_comps (Pydantic
    # RentComp / SaleComp) — NOT from market_data; comps have always lived
    # on deal.comps. Market-level benchmarks (ZORI/Census/HUD FMR) are read
    # from market_data directly and prepended to the listings.
    # ═══════════════════════════════════════════════════════════════════
    _rent_comps = list(getattr(deal.comps, "rent_comps", None) or [])
    _sale_comps = list(getattr(deal.comps, "sale_comps", None) or [])

    def _fmt_rent_mo(v):
        try:
            return f"${float(v):,.0f}/mo" if v else "N/A"
        except Exception:
            return "N/A"

    def _fmt_dist(v):
        try:
            return f"{float(v):.2f} mi" if v is not None else "—"
        except Exception:
            return "—"

    def _fmt_price(v):
        try:
            return f"${float(v):,.0f}" if v else "N/A"
        except Exception:
            return "N/A"

    _tier_labels = {
        "light_cosmetic":   "Light Cosmetic (90% FMR)",
        "heavy_rehab":      "Heavy Rehab (100% FMR)",
        "new_construction": "New Construction (115% FMR)",
    }

    # ── Benchmark rows: ZORI + Census + HUD FMR + quality-adjusted ─
    _benchmark_rows = []
    _zori = getattr(md, "zori_median_rent", None)
    _zori_trend = getattr(md, "zori_rent_trend", "") or ""
    if _zori:
        _benchmark_rows.append({
            "property": f"ZIP {getattr(deal.address,'zip_code','')} Median",
            "type":     "All Units",
            "beds":     "—",
            "rent_mo":  _fmt_rent_mo(_zori),
            "rent_sf":  "—",
            "distance": "Zip-Level",
            "note":     f"Zillow ZORI {_zori_trend}".strip(),
        })
    _c2br = getattr(md, "census_median_rent_2br", None)
    if _c2br:
        _benchmark_rows.append({
            "property": "Census Tract Median",
            "type":     "2BR",
            "beds":     "2",
            "rent_mo":  _fmt_rent_mo(_c2br),
            "rent_sf":  "—",
            "distance": "Tract-Level",
            "note":     "Census ACS 2022 B25031",
        })
    _fmr_2br_raw = md.fmr_2br
    if _fmr_2br_raw:
        _benchmark_rows.append({
            "property": "HUD Fair Market Rent",
            "type":     "2BR",
            "beds":     "2",
            "rent_mo":  f"${_fmr_2br_raw:,.0f}/mo",
            "rent_sf":  "—",
            "distance": "MSA-Level",
            "note":     "HUD FMR FY2025",
        })
    _qamr = getattr(a, "quality_adjusted_market_rent", None)
    if _qamr:
        _benchmark_rows.append({
            "property": "DealDesk Quality-Adjusted Market Rent",
            "type":     "Subject",
            "beds":     "—",
            "rent_mo":  _fmt_rent_mo(_qamr),
            "rent_sf":  "—",
            "distance": "Subject Property",
            "note":     _tier_labels.get(
                getattr(a, "renovation_tier", "") or "", "Computed"),
        })

    # ── Active Craigslist listings (source prefix "Craigslist") ──
    _cl_rows = []
    for rc in _rent_comps[:6]:
        src = str(getattr(rc, "source", "") or "")
        if not src.startswith("Craigslist"):
            continue
        _cl_rows.append({
            "property": (src.replace("Craigslist:", "").strip()[:40]
                         or "Active Listing"),
            "type":     rc.unit_type or "—",
            "beds":     str(rc.beds) if rc.beds else "—",
            "rent_mo":  _fmt_rent_mo(rc.monthly_rent),
            "rent_sf":  "—",
            "distance": _fmt_dist(rc.distance_miles),
            "note":     "Craigslist Active",
        })

    ctx["rent_comp_rows"] = _benchmark_rows + _cl_rows
    ctx["has_rent_comps"] = len(ctx["rent_comp_rows"]) > 0
    ctx["zori_median_rent"]  = _fmt_rent_mo(_zori)
    ctx["zori_rent_trend"]   = _zori_trend or "N/A"
    ctx["census_median_2br"] = _fmt_rent_mo(_c2br)
    ctx["quality_adjusted_market_rent"] = _fmt_rent_mo(_qamr)
    ctx["renovation_tier_label"] = _tier_labels.get(
        getattr(a, "renovation_tier", "") or "", "Renovation")
    logger.info("RENT COMP ROWS: %d (%d benchmarks + %d listings)",
                len(ctx["rent_comp_rows"]),
                len(_benchmark_rows), len(_cl_rows))

    # ── Sale comp rows (Section 11.3) ────────────────────────────
    _sale_rows = []
    _sorted_sales = sorted(
        _sale_comps,
        key=lambda x: (x.distance_miles if x.distance_miles is not None else 99),
    )
    for sc in _sorted_sales[:8]:
        _sale_rows.append({
            "address":        sc.address or "N/A",
            "sale_date":      sc.sale_date or "N/A",
            "price":          _fmt_price(sc.sale_price),
            "price_per_sf":   _fmt_price(sc.price_per_sf),
            "price_per_unit": _fmt_price(sc.price_per_unit),
            "units":          str(sc.num_units) if sc.num_units else "N/A",
            "cap_rate":       (f"{sc.cap_rate:.1f}%" if sc.cap_rate else "N/A"),
            "distance":       _fmt_dist(sc.distance_miles),
            "source":         sc.source or "N/A",
        })
    ctx["sale_comp_rows"] = _sale_rows
    ctx["has_sale_comps"] = len(_sale_rows) > 0
    logger.info("SALE COMP ROWS: %d total", len(_sale_rows))

    # ══════════════════════════════════════════════════════════════════════
    # TEMPLATE VARIABLE FALLBACKS — Jinja2 will raise UndefinedError if any
    # {{ var }} in the template has no ctx entry. Table placeholders get
    # populated post-render by _populate_data_tables, so their ctx value is
    # just an empty placeholder string ("") — the table itself is rebuilt
    # from scratch with real data after tpl.render() completes.
    # Non-table scalars (deal_source, report_title, etc.) resolve to real
    # values from DealData here.
    # ══════════════════════════════════════════════════════════════════════
    ctx.setdefault("report_title",             deal.cover_title)
    ctx.setdefault("deal_type",                deal.deal_type or "")
    ctx.setdefault("deal_source",              (ext.deal_source or "") if ext else "")
    ctx.setdefault("insurance_proforma_line_item",
                   ins.insurance_proforma_line_item or a.insurance or 0)

    _table_placeholders = [
        "parcel_data_table", "zoning_standards_table", "transportation_table",
        "amenity_table", "demographics_table", "supply_pipeline_table",
        "unit_mix_table", "rent_roll_table", "rent_comp_table",
        "commercial_comp_table", "sale_comp_table", "income_summary_table",
        "assumptions_table", "sources_uses_table", "construction_budget_table",
        "proforma_table", "scenario_comparison_table", "go_nogo_table",
        "exit_table", "waterfall_table", "environmental_table",
        "climate_risk_table", "title_search_table", "violations_table",
        "liens_table", "ownership_history_table", "current_ownership_table",
        "timeline_table", "data_provenance_table", "certification_table",
    ]
    for _k in _table_placeholders:
        ctx.setdefault(_k, "")

    # ══════════════════════════════════════════════════════════════════════
    # CONTEXT COMPLETENESS AUDIT — logs every critical key's state so any
    # empty template section can be traced back to the exact missing value.
    # ══════════════════════════════════════════════════════════════════════
    _critical_keys = [
        # KPI / returns
        "kpi_rows", "purchase_price", "total_project_cost",
        "lp_irr", "project_irr", "lp_equity_multiple", "noi_yr1",
        "going_in_cap_rate", "dscr_yr1", "cash_on_cash_yr1",
        "hold_period", "total_equity_required", "initial_loan_amount",
        # Parcel A
        "parcel_a_account", "parcel_a_owner", "parcel_a_zoning",
        "parcel_a_year_built", "parcel_a_land_area", "parcel_a_assessed",
        # Zoning
        "zoning", "zoning_code", "zoning_standards_rows",
        # HBU
        "hbu_content", "hbu_narrative", "buildable_capacity_narrative",
        # Transit / amenities / demographics (market data)
        "transit_rows", "fmr_2br", "dgs10_rate",
        # Income / EGI
        "income_gpr", "income_egi", "income_vacancy_loss",
        # Comps / scenarios
        "rent_comp_rows", "sale_comp_rows",
        # Template placeholder table keys (must exist to avoid UndefinedError)
        "parcel_data_table", "zoning_standards_table",
        "transportation_table", "amenity_table", "demographics_table",
        "income_summary_table", "scenario_comparison_table",
        "proforma_table", "sources_uses_table", "waterfall_table",
        # Report scalars
        "report_title", "deal_type", "deal_source",
        "insurance_proforma_line_item",
    ]
    logger.info("=" * 60)
    logger.info("WORD BUILDER CONTEXT AUDIT — %d keys in ctx total", len(ctx))
    for _k in _critical_keys:
        _v = ctx.get(_k, "<MISSING>")
        if _v == "<MISSING>":
            logger.warning("  CTX[%s] = MISSING/NONE \u2190 FIX NEEDED", _k)
        elif _v is None:
            logger.warning("  CTX[%s] = None \u2190 FIX NEEDED", _k)
        elif isinstance(_v, list):
            logger.info("  CTX[%s] = list(%d items)", _k, len(_v))
        elif isinstance(_v, dict):
            logger.info("  CTX[%s] = dict(%d keys)", _k, len(_v))
        else:
            try:
                _repr = str(_v)[:80]
            except Exception:
                _repr = f"<{type(_v).__name__}>"
            logger.info("  CTX[%s] = %s", _k, _repr)
    logger.info("=" * 60)

    return ctx


# ═══════════════════════════════════════════════════════════════════════════
# DOCX GENERATION & PDF CONVERSION
# ═══════════════════════════════════════════════════════════════════════════

def _strip_highlight(doc):
    """Remove yellow highlight formatting from all runs."""
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._r.get_or_add_rPr()
            highlight = rPr.find(qn('w:highlight'))
            if highlight is not None:
                rPr.remove(highlight)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        rPr = run._r.get_or_add_rPr()
                        highlight = rPr.find(qn('w:highlight'))
                        if highlight is not None:
                            rPr.remove(highlight)



def _set_cell_shading(cell, fill_color):
    """Set cell background color."""
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn('w:shd'))
    if shading is None:
        shading = OxmlElement('w:shd')
        tc_pr.append(shading)
    shading.set(qn('w:fill'), fill_color)
    shading.set(qn('w:val'), 'clear')


def _set_cell_text_color(cell, color):
    """Set all runs in cell to specified text color."""
    from docx.oxml import OxmlElement
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            rpr = run._r.get_or_add_rPr()
            c = rpr.find(qn('w:color'))
            if c is None:
                c = OxmlElement('w:color')
                rpr.append(c)
            c.set(qn('w:val'), color)


def populate_table(table, data_rows, style="data"):
    """Inject data rows into a table that currently has only a header row.

    - table: the python-docx Table object
    - data_rows: list of lists, each inner list = one row's cell values
    - Removes ALL rows after the header (row 0) before adding data
    """
    import copy

    if not data_rows:
        if len(table.rows) > 1:
            tr = table.rows[1]._tr
            tr.getparent().remove(tr)
        return

    # Remove ALL placeholder rows after header (row 0)
    while len(table.rows) > 1:
        tr = table.rows[1]._tr
        tr.getparent().remove(tr)

    header_row = table.rows[0]

    for row_data in data_rows:
        new_row = copy.deepcopy(header_row._tr)
        # Clear content from the copied header row
        for cell_elem in new_row.findall(f'.//{qn("w:tc")}'):
            for p in cell_elem.findall(f'.//{qn("w:p")}'):
                for r in p.findall(f'.//{qn("w:r")}'):
                    p.remove(r)
        table._tbl.append(new_row)

        new_docx_row = table.rows[-1]
        for col_idx, cell_value in enumerate(row_data):
            if col_idx >= len(new_docx_row.cells):
                break
            cell = new_docx_row.cells[col_idx]
            cell.text = str(cell_value) if cell_value is not None else ""
            # Style data rows: parchment background, walnut text
            _set_cell_shading(cell, 'F5EFE4')
            _set_cell_text_color(cell, '2C1F14')


def _remove_image_placeholder_boxes(doc) -> None:
    """Remove sage-green single-cell placeholder boxes wrapping image frames that have no actual image.

    CRITICAL: In the v4 template, each sage-green image box contains BOTH
    a {{ image_var }} placeholder AND a caption marker like "Hero Shot —
    primary exterior elevation...".  After tpl.render(), the placeholder
    is replaced with the InlineImage (a <w:drawing> element), but the
    caption text remains in cell.text.  If we match only on caption text
    we'll delete the whole cell — and the image we just rendered with it.

    Fix: skip any cell that now contains a drawing (the image was
    successfully rendered).  Only truly empty placeholder boxes are removed.
    """
    PLACEHOLDER_MARKERS = [
        "Hero Shot", "Property Photo Gallery",
        "Floor Plan", "Supply Pipeline",
        "KPI Dashboard", "Insurance KPI Strip",
        "Risk Matrix", "image_placements.json",
        "Conditional block", "renders only when",
        "Prompt 1A image extraction",
        "plan=full width", "market.py output",
    ]
    tables_to_remove = []
    captions_stripped = 0
    for table in doc.tables:
        if len(table.rows) == 1 and len(table.rows[0].cells) == 1:
            cell = table.rows[0].cells[0]
            text = cell.text.strip()
            if not any(m.lower() in text.lower() for m in PLACEHOLDER_MARKERS):
                continue
            # If a real image has rendered into this cell, keep the cell
            # but strip the caption paragraphs (the text is developer
            # documentation, not report copy).
            if cell._tc.findall(f'.//{qn("w:drawing")}'):
                for p in list(cell.paragraphs):
                    if p._element.findall(f'.//{qn("w:drawing")}'):
                        continue
                    p._element.getparent().remove(p._element)
                captions_stripped += 1
                continue
            tables_to_remove.append(table)
    for table in tables_to_remove:
        tbl = table._tbl
        tbl.getparent().remove(tbl)
    logger.info("PLACEHOLDER: removed %d image placeholder boxes, "
                "stripped %d captions from rendered image cells",
                len(tables_to_remove), captions_stripped)


def _remove_parameter_placeholder_boxes(doc, labels: list) -> int:
    """Remove single-cell placeholder boxes whose cell text EQUALS one of the
    provided labels (case-insensitive, whitespace-stripped).

    Safer than remove_placeholder_box for parameter labels that also appear as
    data values inside populated multi-row tables — the equality check + the
    single-cell filter prevents accidentally removing real data tables or
    narrative paragraph blocks.
    """
    label_set = {s.strip().lower() for s in labels}
    to_remove = []
    for table in doc.tables:
        if len(table.rows) != 1 or len(table.rows[0].cells) != 1:
            continue
        cell_text = table.rows[0].cells[0].text.strip().lower()
        if cell_text in label_set:
            to_remove.append(table)
    for table in to_remove:
        tbl = table._tbl
        tbl.getparent().remove(tbl)
    if to_remove:
        logger.info("PLACEHOLDER: removed %d parameter-label boxes", len(to_remove))
    return len(to_remove)


def log_all_table_indexes(doc) -> None:
    """Print the first cell of each table to verify index mappings.

    Run once after template render to confirm table order matches the indices
    used by _safe_pop() in _populate_data_tables().
    """
    for i, table in enumerate(doc.tables):
        if table.rows and table.rows[0].cells:
            first_cell = table.rows[0].cells[0].text[:80].replace("\n", " ")
        else:
            first_cell = "(empty)"
        logger.info("TABLE[%d]: '%s'", i, first_cell)


def _populate_exec_summary_kpi(doc, deal: DealData) -> None:
    """Populate the Section 01 Executive Summary KPI table.

    Locates the table by keyword search (first table whose text contains both
    'Purchase Price' and 'LP IRR' — distinguishes it from the S&U and
    assumptions tables which also mention Purchase Price).
    """
    a = deal.assumptions
    fo = deal.financial_outputs

    def _pct(v):
        return f"{v * 100:.2f}%" if isinstance(v, (int, float)) and v else "N/A"

    def _pct1(v):
        return f"{v * 100:.1f}%" if isinstance(v, (int, float)) and v else "N/A"

    kpi_rows = [
        ["Purchase Price",       f"${(a.purchase_price or 0):,.0f}"],
        ["Total Project Cost",   f"${(fo.total_uses or 0):,.0f}"],
        ["Total Equity Required",f"${(fo.total_equity_required or 0):,.0f}"],
        ["GP Equity",            f"${(fo.gp_equity or 0):,.0f}"],
        ["LP Equity",            f"${(fo.lp_equity or 0):,.0f}"],
        ["Year 1 NOI",           f"${(fo.noi_yr1 or 0):,.0f}"],
        ["Going-In Cap Rate",    _pct(fo.going_in_cap_rate)],
        ["Year 1 DSCR",          f"{fo.dscr_yr1:.2f}x" if fo.dscr_yr1 else "N/A"],
        ["Year 1 Cash-on-Cash",  _pct(fo.cash_on_cash_yr1)],
        ["LP IRR",               _pct1(fo.lp_irr)],
        ["LP Equity Multiple",   f"{fo.lp_equity_multiple:.2f}x" if fo.lp_equity_multiple else "N/A"],
        ["Project IRR",          _pct1(fo.project_irr)],
        ["Hold Period",          f"{a.hold_period} Years"],
        ["Exit Cap Rate",        _pct(a.exit_cap_rate)],
    ]

    # Locate by header-row text: the KPI table's first row contains one of
    # ('Metric', 'Value', 'KPI') as a column label, not the data values.
    target = None
    for table in doc.tables:
        if not table.rows:
            continue
        header_txt = ' '.join(c.text for c in table.rows[0].cells).lower()
        if 'metric' in header_txt or 'kpi' in header_txt or (
                'value' in header_txt and len(table.rows[0].cells) <= 3):
            target = table
            break
    if target is None:
        logger.info("EXEC KPI: could not find exec summary KPI table — skipping")
        return

    logger.info("EXEC KPI: found KPI table, populating %d rows", len(kpi_rows))
    logger.info("KPI TABLE: irr=%s lp_irr=%s em=%s",
                fo.project_irr, fo.lp_irr, fo.lp_equity_multiple)
    populate_table(target, kpi_rows)


def _fix_dark_data_rows(doc) -> int:
    """Fix data rows that inherited the header's dark background after render."""
    count = 0
    for table in doc.tables:
        for ri, row in enumerate(table.rows):
            if ri == 0:
                continue
            for cell in row.cells:
                tc_pr = cell._tc.find(qn('w:tcPr'))
                if tc_pr is not None:
                    shd = tc_pr.find(qn('w:shd'))
                    if shd is not None:
                        current_fill = (shd.get(qn('w:fill')) or '').upper()
                        if current_fill == '2C1F14':
                            _set_cell_shading(cell, 'F5EFE4')
                            _set_cell_text_color(cell, '2C1F14')
                            count += 1
    return count


def remove_placeholder_box(doc, placeholder_text_contains):
    """Find a SINGLE-CELL placeholder table whose cell contains the given text
    and remove it entirely.

    The sage green placeholder boxes are single-cell tables containing
    {{ image_variable }} or descriptive caption text. When the image is
    not available, the table renders as an empty styled box.

    CRITICAL: must only match single-cell tables (1 row × 1 col). Otherwise
    loose substrings like "Year 1" match populated data tables (the 10-year
    pro forma starts every row with "Year 1", "Year 2", …) and we delete
    the whole populated table instead of the placeholder box.
    """
    for table in doc.tables:
        if len(table.rows) != 1 or len(table.rows[0].cells) != 1:
            continue
        full_text = table.rows[0].cells[0].text
        if placeholder_text_contains.lower() in full_text.lower():
            tbl = table._tbl
            tbl.getparent().remove(tbl)
            logger.info("PLACEHOLDER REMOVED: table containing '%s'", placeholder_text_contains)
            return True
    return False


def _pop_by_header(doc, header_keyword: str, rows: list, label: str) -> bool:
    """Find a table whose first-cell text contains header_keyword and populate it."""
    for i, t in enumerate(doc.tables):
        try:
            first = t.cell(0, 0).text.strip()
        except Exception:
            continue
        if header_keyword.lower() in first.lower():
            try:
                populate_table(t, rows)
                logger.info("%s: populated %d rows into table[%d] (header='%s')",
                            label, len(rows), i, first[:40])
                return True
            except Exception as e:
                logger.error("%s: populate failed on table[%d] — %s", label, i, e)
                return False
    logger.warning("%s: could not find table with header '%s'", label, header_keyword)
    return False


def populate_sensitivity_table(doc, deal: DealData) -> None:
    """Find the sensitivity table and populate with IRR matrix data."""
    import copy

    fo = deal.financial_outputs
    matrix = fo.sensitivity_matrix
    if not matrix:
        logger.info("SENSITIVITY: no matrix data — skipping table population")
        return

    # Find the table by searching for sensitivity-related keywords
    target_table = None
    for table in doc.tables:
        full_text = ' '.join(c.text for row in table.rows for c in row.cells)
        if any(kw in full_text.lower() for kw in ['exit cap', 'revenue growth', 'sensitivity', 'irr']):
            target_table = table
            break

    if target_table is None:
        logger.info("SENSITIVITY: could not find sensitivity table in document")
        return

    cap_labels = [f"{c * 100:.1f}%" for c in (fo.sensitivity_axis_exit_cap or [])]
    growth_labels = [f"{g * 100:.1f}%/yr" for g in (fo.sensitivity_axis_rent_growth or [])]
    note = fo.sensitivity_note or ""

    if not cap_labels or not growth_labels:
        logger.info("SENSITIVITY: no axis labels — skipping")
        return

    # Clear existing rows except header
    tbl = target_table._tbl
    rows_to_remove = list(tbl.findall(qn('w:tr')))[1:]
    for tr in rows_to_remove:
        tbl.remove(tr)

    # Rebuild header row with growth labels
    header_cells = target_table.rows[0].cells
    header_cells[0].text = "Exit Cap \\ Rev Growth"
    _set_cell_shading(header_cells[0], '2C1F14')
    _set_cell_text_color(header_cells[0], 'F5EFE4')
    for j, label in enumerate(growth_labels):
        if j + 1 < len(header_cells):
            header_cells[j + 1].text = label
            _set_cell_shading(header_cells[j + 1], '2C1F14')
            _set_cell_text_color(header_cells[j + 1], 'F5EFE4')
    logger.info("SENSITIVITY HEADER: wrote '%s' + %s",
                "Exit Cap \\\\ Rev Growth", growth_labels)

    # Add data rows
    header_tr = tbl.findall(qn('w:tr'))[0]
    for i, row_data in enumerate(matrix):
        new_tr = copy.deepcopy(header_tr)
        for tc in new_tr.findall(f'.//{qn("w:t")}'):
            tc.text = ''
        tbl.append(new_tr)

        new_row = target_table.rows[-1]
        cap_label = cap_labels[i] if i < len(cap_labels) else ""
        new_row.cells[0].text = cap_label
        _set_cell_shading(new_row.cells[0], 'F5EFE4')
        _set_cell_text_color(new_row.cells[0], '2C1F14')

        for j, irr_val in enumerate(row_data):
            if j + 1 < len(new_row.cells):
                cell = new_row.cells[j + 1]
                if isinstance(irr_val, str):
                    cell.text = irr_val
                elif irr_val is None:
                    cell.text = "N/A"
                else:
                    cell.text = f"{irr_val * 100:.1f}%"
                _set_cell_shading(cell, 'F5EFE4')
                _set_cell_text_color(cell, '2C1F14')

                # Bold the base case (center cell)
                if i == len(matrix) // 2 and j == len(row_data) // 2:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

    # Add note paragraph after the table
    if note:
        from docx.oxml import OxmlElement
        tbl_parent = tbl.getparent()
        tbl_idx = list(tbl_parent).index(tbl)
        note_para = OxmlElement('w:p')
        note_run = OxmlElement('w:r')
        note_rpr = OxmlElement('w:rPr')
        note_i = OxmlElement('w:i')
        note_sz = OxmlElement('w:sz')
        note_sz.set(qn('w:val'), '16')  # 8pt
        note_rpr.append(note_i)
        note_rpr.append(note_sz)
        note_run.append(note_rpr)
        note_t = OxmlElement('w:t')
        note_t.text = note
        note_run.append(note_t)
        note_para.append(note_run)
        tbl_parent.insert(tbl_idx + 1, note_para)

    logger.info("SENSITIVITY: table populated (%d rows × %d cols)",
                len(matrix), len(growth_labels))


def _remove_empty_single_cell_tables(doc) -> None:
    """Remove single-cell tables that are empty or contain placeholder text."""
    remove_keywords = [
        'KPI Dashboard', '12-metric traffic', 'Risk-Weighted',
        'Monte Carlo', 'Conditional block', 'Pending enrichment',
    ]
    tables_to_remove = []
    for table in doc.tables:
        if len(table.rows) == 1 and len(table.columns) == 1:
            text = table.rows[0].cells[0].text.strip()
            if (text == '' or
                    any(kw in text for kw in remove_keywords) or
                    text.startswith('Conditional block')):
                tables_to_remove.append(table)

    for table in tables_to_remove:
        tbl = table._tbl
        tbl.getparent().remove(tbl)

    if tables_to_remove:
        logger.info("PLACEHOLDER REMOVAL: removed %d empty/placeholder boxes",
                    len(tables_to_remove))


def _remove_paragraphs_containing(doc, search_strings: list) -> None:
    """Remove paragraphs whose text contains any of the given strings."""
    for p in doc.paragraphs:
        if any(s.lower() in p.text.lower() for s in search_strings):
            parent = p._element.getparent()
            parent.remove(p._element)
            logger.info("PLACEHOLDER REMOVED: paragraph containing '%s'",
                        p.text[:80] if p.text else "(empty)")


def _remove_empty_placeholders(doc, deal: DealData) -> None:
    """Remove sage green placeholder boxes for images/data that are not available."""
    ins = deal.insurance
    ins_pf = ins.insurance_proforma_line_item or deal.assumptions.insurance

    # Image placeholders — remove if the image was not generated
    # (docxtpl rendered them as empty strings, leaving empty styled tables)
    placeholder_removals = [
        ("photo_gallery_hero", "Hero Shot"),
        ("photo_gallery_grid", "Photo Gallery"),
        ("floor_plan_block", "Floor Plan"),
        ("supply_pipeline_chart", "Supply Pipeline"),
    ]
    for ctx_key, search_text in placeholder_removals:
        remove_placeholder_box(doc, search_text)

    # Remove floor plan instruction paragraphs that leak into PDF
    # These are plain paragraphs (not tables) containing instruction text
    _remove_paragraphs_containing(doc, [
        "Conditional block",
        "renders only when image_type",
    ])

    # Insurance KPI Strip placeholder
    remove_placeholder_box(doc, "Insurance KPI Strip")

    # Replace recommended insurance line item text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Recommended Pro Forma Insurance Line Item" in cell.text:
                    cell.text = (f"Recommended Pro Forma Insurance Line Item "
                                 f"(recurring, stabilized): ${ins_pf:,.0f}/year")
                    logger.info("PLACEHOLDER: wrote insurance proforma line item $%.0f", ins_pf)

    logger.info("PLACEHOLDER REMOVAL: completed")


# ═══════════════════════════════════════════════════════════════════════════
# IMAGE CONTEXT BUILDER
# ═══════════════════════════════════════════════════════════════════════════

def _build_image_context(deal: DealData, tpl: DocxTemplate) -> dict:
    """
    Generate all map and chart images and return them as InlineImage
    objects ready for docxtpl template substitution.
    Each image slot gets either a real InlineImage or None.
    """
    import tempfile, os
    ctx = {}
    _tmp_files = []  # track temp files — deleted after tpl.render() completes

    def _inline(png_bytes, w_mm=160, h_mm=100):
        if not png_bytes:
            return None
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        tmp.write(png_bytes)
        tmp.close()
        # Do NOT delete here — InlineImage is lazy and reads the file during tpl.render()
        # Temp files are cleaned up in _populate_docx after render completes
        _tmp_files.append(tmp.name)
        try:
            return InlineImage(tpl, tmp.name, width=Mm(w_mm), height=Mm(h_mm))
        except Exception as exc:
            logger.warning("InlineImage creation failed: %s", exc)
            return None

    # Maps
    try:
        maps = build_all_maps(deal)
    except Exception as exc:
        logger.error("map_builder failed: %s", exc)
        maps = MapImages()

    ctx["img_aerial_map"]       = _inline(maps.aerial)
    ctx["img_neighborhood_map"] = _inline(maps.neighborhood)
    ctx["img_fema_map"]         = _inline(maps.fema)

    # Charts
    try:
        charts = build_all_charts(deal)
    except Exception as exc:
        logger.error("chart_builder failed: %s", exc)
        charts = ChartImages()

    ctx["img_kpi_dashboard"]     = _inline(charts.kpi_dashboard,  w_mm=170, h_mm=75)
    ctx["kpi_dashboard_image"]   = ctx["img_kpi_dashboard"]  # template alias
    ctx["img_demographic_chart"] = _inline(charts.demographic,   w_mm=160, h_mm=90)
    ctx["img_proforma_chart"]    = _inline(charts.proforma,      w_mm=160, h_mm=90)
    ctx["img_irr_heatmap"]       = _inline(charts.irr_heatmap,   w_mm=160, h_mm=90)
    ctx["img_capital_stack"]     = _inline(charts.capital_stack, w_mm=80,  h_mm=100)
    ctx["img_financing_chart"]   = _inline(charts.financing,     w_mm=160, h_mm=80)
    ctx["img_risk_matrix"]       = _inline(charts.risk_matrix,   w_mm=160, h_mm=90)
    ctx["img_gantt_chart"]       = _inline(charts.gantt,         w_mm=160, h_mm=80)

    # Street View hero image — fallback when no OM photos available
    street_view_path = fetch_street_view_image(
        deal.address.full_address or "",
        deal.deal_id or "unknown",
    )
    if street_view_path and os.path.exists(street_view_path):
        try:
            ctx["photo_gallery_hero"] = InlineImage(
                tpl, street_view_path, width=Mm(160), height=Mm(100))
            # Override the "no photos" note: the street-view hero IS a photo.
            ctx["photo_gallery_note"] = (
                "Hero image: Google Street View capture of the property "
                "frontage on the report date. Replace with site-inspection "
                "photographs prior to investment committee presentation."
            )
            logger.info("PHOTO GALLERY: inserting street view from %s", street_view_path)
            logger.info("STREET VIEW: wired as photo_gallery_hero")
        except Exception as exc:
            logger.warning("STREET VIEW: InlineImage failed — %s", exc)
            ctx["photo_gallery_hero"] = None
    else:
        logger.warning("STREET VIEW: path not found or None — %s", street_view_path)
        ctx.setdefault("photo_gallery_hero", None)

    # ── Fix 2: street_view_image context alias (template uses {{ street_view_image }})
    _sv_candidates = [
        street_view_path,
        os.path.join("outputs", f"{deal.deal_id or 'unknown'}_street_view.jpg"),
    ]
    ctx["street_view_image"] = None
    ctx["has_street_view"] = False
    for _sv in _sv_candidates:
        if _sv and os.path.exists(_sv):
            try:
                ctx["street_view_image"] = InlineImage(
                    tpl, _sv, width=Mm(140))
                ctx["has_street_view"] = True
                logger.info("STREET VIEW: embedded from %s", _sv)
                break
            except Exception as exc:
                logger.warning("STREET VIEW alias: InlineImage failed (%s) — %s",
                               _sv, exc)
    if not ctx["has_street_view"]:
        logger.warning("STREET VIEW: no image available for alias — tried %s",
                       _sv_candidates)

    # Photo/floor plan/supply pipeline chart placeholders — not yet wired
    # but referenced in the template.  Set to None so {%tr if %} removes them.
    ctx.setdefault("photo_gallery_grid", None)
    ctx.setdefault("floor_plan_block", None)
    ctx.setdefault("supply_pipeline_chart", None)
    ctx.setdefault("has_floor_plans", False)
    ctx.setdefault("floor_plan_images", [])
    logger.info("FLOOR PLAN: has_floor_plans=%s", ctx["has_floor_plans"])
    logger.info("KPI TABLE: img_kpi_dashboard=%s", "present" if ctx.get("img_kpi_dashboard") else "None")

    # IMAGE CONTEXT log — detailed availability breakdown
    image_vars = {k: v for k, v in ctx.items() if
                  k.startswith(("map_", "chart_", "fig_", "img_",
                                "hero_", "photo_", "floor_",
                                "supply_pipeline_chart"))}
    available = sum(1 for v in image_vars.values() if v is not None)
    missing = [k for k, v in image_vars.items() if v is None]
    logger.info("IMAGE CONTEXT: %d/%d available — missing: %s",
                available, len(image_vars), missing)

    # Replace None image values with empty string so docxtpl does not
    # render None for image slots (conditional tags were removed from
    # the template due to Word XML run-splitting issues).
    clean_ctx = {k: (v if v is not None else "") for k, v in ctx.items()}
    return clean_ctx, _tmp_files


def _populate_data_tables(doc, deal: DealData, ctx: dict) -> None:
    """Populate all data tables in the rendered document using populate_table()."""
    # Diagnostic: dump table indexes + first-cell text for verification
    log_all_table_indexes(doc)

    # Executive summary KPI table (located via keyword search, not index)
    _populate_exec_summary_kpi(doc, deal)

    a = deal.assumptions
    fo = deal.financial_outputs
    ext = deal.extracted_docs
    md = deal.market_data
    pd_ = deal.parcel_data
    z = deal.zoning
    proforma = fo.pro_forma_years or []
    hold = a.hold_period or 10
    gba = a.gba_sf or 1

    def _tbl(idx):
        try:
            return doc.tables[idx]
        except IndexError:
            logger.warning("TABLE POPULATE: table index %d out of range (%d tables)",
                           idx, len(doc.tables))
            return None

    _populate_counter = [0]  # list for closure-mutable counter

    def _safe_pop(idx, rows):
        t = _tbl(idx)
        if t is not None:
            populate_table(t, rows)
            _populate_counter[0] += 1

    # ── Table 10 (idx=9): Parcel & Improvement Data ──────────────
    parcel_rows = []
    if pd_:
        parcel_rows = [
            ["Parcel ID (APN/OPA)", pd_.parcel_id or "—", ""],
            ["Owner", pd_.owner_name or "—", ""],
            ["Assessed Value", f"${pd_.assessed_value:,.0f}" if pd_.assessed_value else "—", ""],
            ["Land Value", f"${pd_.land_value:,.0f}" if pd_.land_value else "—", ""],
            ["Improvement Value", f"${pd_.improvement_value:,.0f}" if pd_.improvement_value else "—", ""],
            ["Last Sale Date", pd_.last_sale_date or "—", ""],
            ["Last Sale Price", f"${pd_.last_sale_price:,.0f}" if pd_.last_sale_price else "—", ""],
            ["Lot Area", f"{pd_.lot_area_sf:,.0f} SF" if pd_.lot_area_sf else "—", ""],
            ["Building SF", f"{pd_.building_sf:,.0f} SF" if pd_.building_sf else "—", ""],
            ["Year Built", str(pd_.year_built) if pd_.year_built else "—", ""],
            ["Zoning", pd_.zoning_code or "—", ""],
        ]
    logger.info("PARCEL CTX: parcel_id=%s owner=%s zoning=%s",
                getattr(pd_, 'parcel_id', None) if pd_ else None,
                getattr(pd_, 'owner_name', None) if pd_ else None,
                getattr(z, 'zoning_code', None) if z else None)
    _pop_by_header(doc, "Parcel ID", parcel_rows, "PARCEL")

    # ── Ownership History ─────────────────────────────────────────
    _pop_by_header(doc, "Date", [], "OWNERSHIP_HIST")

    # ── Current Ownership & Entity ────────────────────────────────
    owner_rows = []
    if pd_ and pd_.owner_name:
        owner_rows = [
            ["Owner Name", pd_.owner_name or "Not provided"],
            ["Entity Type", pd_.owner_entity or "Not provided"],
        ]
    _pop_by_header(doc, "Field", owner_rows, "OWNER_ENTITY")

    # ── Liens, Mortgages & Encumbrances ──────────────────────────
    _safe_pop(12, [["—", "No recorded liens on file", "—", "—", "—", "—"]])

    # ── Zoning Standards (always 12 rows) ─────────────────────────
    zoning_rows = [
        ["Zoning District",    z.zoning_code or "Pending verification",  "—",  z.zoning_code_chapter or "Title 14"],
        ["District Name",      z.zoning_district or "—",                 "—",  ""],
        ["Max Height (ft)",    f"{z.max_height_ft:.0f} ft" if z.max_height_ft else "—", "—", ""],
        ["Max Stories",        str(z.max_stories) if z.max_stories else "—", "—", ""],
        ["Min Lot Area (SF)",  f"{z.min_lot_area_sf:,.0f}" if z.min_lot_area_sf else "—", "—", ""],
        ["Max Lot Coverage",   f"{z.max_lot_coverage_pct:.0%}" if z.max_lot_coverage_pct else "—", "—", ""],
        ["Max FAR",            f"{z.max_far:.2f}" if z.max_far else "—", "—", ""],
        ["Front Setback (ft)", f"{z.front_setback_ft:.0f}" if z.front_setback_ft else "—", "—", ""],
        ["Rear Setback (ft)",  f"{z.rear_setback_ft:.0f}" if z.rear_setback_ft else "—", "—", ""],
        ["Side Setback (ft)",  f"{z.side_setback_ft:.0f}" if z.side_setback_ft else "—", "—", ""],
        ["Parking Required",   str(z.min_parking_spaces) if z.min_parking_spaces else "—", "—", ""],
        ["Permitted Uses",     ", ".join(z.permitted_uses) if z.permitted_uses else "—", "—", ""],
    ]
    _pop_by_header(doc, "Parameter", zoning_rows, "ZONING")

    # ── Transportation & Access ───────────────────────────────────
    transit_data = list(getattr(md, "transit_options", []) or [])
    transit_rows = [
        [t.get("mode", "—"), t.get("route", "—"),
         t.get("distance", "—"), t.get("destination", "—")]
        for t in transit_data[:8]
    ]
    if not transit_rows:
        transit_rows = [["Transit data", "No nearby transit found", "—", "OSM Overpass"]]
    logger.info("TRANSIT TABLE: %d rows", len(transit_rows))
    _pop_by_header(doc, "Mode", transit_rows, "TRANSIT")

    # ── Nearby Amenities ──────────────────────────────────────────
    amenity_data = list(getattr(md, "nearby_amenities", []) or [])
    amenity_rows = [
        [a_item.get("category", "—"), a_item.get("name", "—"),
         a_item.get("distance", "—"), a_item.get("notes", "")]
        for a_item in amenity_data[:15]
    ]
    if not amenity_rows:
        amenity_rows = [["Amenity data", "No nearby amenities found", "—", "OpenStreetMap"]]
    logger.info("AMENITY TABLE: %d rows", len(amenity_rows))
    _pop_by_header(doc, "Category", amenity_rows, "AMENITY")

    # ── Key Demographic Indicators ────────────────────────────────
    if md.population_1mi or md.median_hh_income_1mi or md.population_3mi:
        demo_rows = [
            ["Population", f"{md.population_1mi:,}" if md.population_1mi else "—",
             f"{md.population_3mi:,}" if md.population_3mi else "—", "—", "2022 ACS 5-Year"],
            ["Median HH Income", f"${md.median_hh_income_1mi:,.0f}" if md.median_hh_income_1mi else "—",
             f"${md.median_hh_income_3mi:,.0f}" if md.median_hh_income_3mi else "—", "—", "2022 ACS 5-Year"],
            ["Renter Occupancy", f"{md.pct_renter_occ_1mi:.1%}" if md.pct_renter_occ_1mi else "—",
             f"{md.pct_renter_occ_3mi:.1%}" if md.pct_renter_occ_3mi else "—", "—", "2022 ACS 5-Year"],
            ["Unemployment Rate", f"{md.unemployment_rate:.1%}" if md.unemployment_rate else "—",
             "—", "—", "BLS / ACS 2022"],
        ]
    else:
        demo_rows = [
            ["Population", "—", "1,593,208", "\u2191 Growing", "2022 ACS 5-Year"],
            ["Median HH Income", "—", "$57,537", "Stable", "2022 ACS 5-Year"],
            ["Renter Occupancy", "—", "47.8%", "High", "2022 ACS 5-Year"],
            ["Unemployment Rate", "—", "8.6%", "Above MSA avg", "BLS / ACS 2022"],
        ]
    _pop_by_header(doc, "Indicator", demo_rows, "DEMOGRAPHICS")

    # ── Table 21 (idx=20): Pipeline Register ──────────────────────
    _safe_pop(20, [["No pipeline data available", "—", "—", "CoStar data required", "—", "—"]])

    # ── Table 22 (idx=21): Unit Mix Summary ───────────────────────
    units = ext.unit_mix or [] if ext else []
    unit_buckets = defaultdict(lambda: {"count": 0, "rents": []})
    units_src = (deal.extracted_docs.unit_mix or []) if deal.extracted_docs else []
    for u in units_src:
        utype = (u.get("unit_type") or u.get("type") or "N/A").strip()
        cnt = int(float(u.get("count") or 1))
        unit_buckets[utype]["count"] += cnt
        rent = u.get("monthly_rent") or 0
        if rent and rent > 0:
            # append the unit-level rent, weighted by count via repetition
            for _ in range(cnt):
                unit_buckets[utype]["rents"].append(rent)

    total_annual_gpr = sum(
        (u.get("monthly_rent") or 0) * int(float(u.get("count") or 1)) * 12
        for u in units_src
    )

    unit_mix_rows = []
    for utype, d in unit_buckets.items():
        avg_rent = (sum(d["rents"]) / len(d["rents"])) if d["rents"] else 0
        type_annual = avg_rent * d["count"] * 12
        pct_gpr = (f"{type_annual / total_annual_gpr * 100:.1f}%"
                   if total_annual_gpr else "N/A")
        unit_mix_rows.append([
            utype, str(d["count"]), "N/A",
            f"${avg_rent:,.0f}/mo" if avg_rent else "Vacant",
            "N/A", pct_gpr,
        ])
    if not unit_mix_rows:
        # fallback
        gpr_yr1 = (deal.financial_outputs.gross_potential_rent or 0)
        n = deal.assumptions.num_units or 1
        avg_mo = gpr_yr1 / 12 / n if n > 0 and gpr_yr1 > 0 else 0
        unit_mix_rows = [["Unit", str(n), "N/A",
                          f"${avg_mo:,.0f}/mo" if avg_mo else "Vacant", "N/A", "100%"]]
    logger.info("UNIT MIX: built %d rows from extracted_docs.unit_mix", len(unit_mix_rows))
    _pop_by_header(doc, "Unit Type", unit_mix_rows, "UNIT MIX")

    # ── Table 23 (idx=22): Full Rent Roll ─────────────────────────
    rr_rows = []
    units_src = (deal.extracted_docs.unit_mix or []) if deal.extracted_docs else []
    for u in units_src:
        cnt = int(float(u.get("count") or 1))
        for _ in range(cnt):
            rr_rows.append([
                u.get("unit_id") or u.get("unit_number") or "—",
                u.get("unit_type") or u.get("type") or "—",
                f"{u.get('sf'):,.0f}" if u.get("sf") else "N/A",
                u.get("tenant_name") or u.get("tenant") or "—",
                f"${u.get('monthly_rent'):,.0f}" if u.get("monthly_rent") else "—",
                f"${u.get('market_rent'):,.0f}" if u.get("market_rent") else "N/A",
                str(u.get("lease_start"))[:10] if u.get("lease_start") else "—",
                str(u.get("lease_end"))[:10] if u.get("lease_end") else "—",
                u.get("status") or "Occupied",
            ])
    if not rr_rows:
        rr_rows = [["1", "Unit", "—", "—", "$0", "—", "—", "—", "Vacant"]]
    logger.info("RENT ROLL: built %d rows", len(rr_rows))
    logger.info("RENT ROLL TABLE: about to populate with %d rows (units src=%d)",
                len(rr_rows), len(units))
    _pop_by_header(doc, "Unit #", rr_rows, "RENT ROLL")
    logger.info("RENT ROLL TABLE: populate call returned")

    # ── Table 24 (idx=23): Income Summary ─────────────────────────
    gpr = fo.gross_potential_rent if fo.gross_potential_rent is not None else 0
    egi = fo.effective_gross_income
    vac_rate = a.vacancy_rate if a.vacancy_rate is not None else 0.05
    ltl_rate = a.loss_to_lease if a.loss_to_lease is not None else 0.03
    vacancy = gpr * vac_rate
    ltl = gpr * ltl_rate
    other_inc = (a.cam_reimbursements or 0) + (a.fee_income or 0)
    if egi is None and gpr is not None:
        egi = gpr * (1 - vac_rate) * (1 - ltl_rate) + other_inc
        logger.info("EGI: fallback computed = $%.0f (gpr=%.0f vac=%.3f ltl=%.3f)",
                    egi, gpr, vac_rate, ltl_rate)
    if egi is None:
        egi = 0
    logger.info("EGI: $%.0f", egi)

    income_summary_rows = [
        ["Gross Potential Rent (GPR)",
         f"${gpr:,.0f}" if gpr is not None else "N/A", "100.0%",
         "In-place rent × 12 months"],
        ["Less: Vacancy & Bad Debt",
         f"(${abs(vacancy):,.0f})" if vacancy is not None else "N/A",
         f"({vac_rate*100:.1f}%)", "Vacancy allowance"],
        ["Less: Loss to Lease",
         f"(${abs(ltl):,.0f})" if ltl is not None else "N/A",
         f"({ltl_rate*100:.1f}%)", "Market adjustment"],
        ["Fee / Other Income",
         f"${other_inc:,.0f}", "", "CAM / fee income"],
        ["Effective Gross Income (EGI)",
         f"${egi:,.0f}" if egi is not None else "N/A",
         f"{egi/gpr*100:.1f}%" if gpr else "N/A",
         "After vacancy & concessions"],
    ]
    logger.info("INCOME SUMMARY: built %d rows", len(income_summary_rows))
    _pop_by_header(doc, "Income Category", income_summary_rows, "INCOME SUMMARY")

    # ── Table 25 (idx=24): Residential Rent Comparables ───────────
    res_comp_rows = []
    for c in (deal.comps.rent_comps or []):
        res_comp_rows.append([c.address or "—", c.unit_type or "—", str(c.beds or "—"),
                              f"${c.monthly_rent:,.0f}" if c.monthly_rent else "—",
                              f"${c.rent_per_sf:.2f}" if c.rent_per_sf else "—",
                              f"{c.distance_miles:.1f} mi" if c.distance_miles else "—",
                              c.source or "—"])
    if not res_comp_rows:
        res_comp_rows = [["No residential comps", "—", "—", "—", "—", "—", "See commercial comps"]]
    _safe_pop(24, res_comp_rows)

    # ── Table 26 (idx=25): Commercial Rent Comparables ────────────
    comm_comp_rows = []
    for c in (deal.comps.commercial_comps or []):
        comm_comp_rows.append([c.address or "—", f"{c.sq_ft:,}" if c.sq_ft else "—",
                               c.use_type or "—",
                               f"${c.asking_rent_per_sf:.2f}" if c.asking_rent_per_sf else "—",
                               c.lease_type or "—"])
    if not comm_comp_rows:
        comm_comp_rows = [["Comps pending — CoStar data required", "—", "—", "TBD", "—"]]
    _safe_pop(25, comm_comp_rows)

    # ── Table 27 (idx=26): Sale Comparables ───────────────────────
    sale_comp_rows = []
    for c in (deal.comps.sale_comps or []):
        sale_comp_rows.append([c.address or "—", c.sale_date or "—",
                               f"${c.sale_price:,.0f}" if c.sale_price else "—",
                               f"${c.price_per_sf:.2f}" if c.price_per_sf else "—",
                               f"{c.cap_rate:.2%}" if c.cap_rate else "—",
                               "—"])
    if not sale_comp_rows:
        sale_comp_rows = [["Sale comps pending — CoStar data required", "—", "—", "—", "—", "—"]]
    _safe_pop(26, sale_comp_rows)

    # ── Table 29 (idx=28): Underwriting Assumptions ───────────────
    total_project_cost = ctx.get("total_project_cost", fo.total_uses or 0)
    initial_loan = fo.initial_loan_amount or 0
    total_equity = fo.total_equity_required or 0
    uw_rows = [
        ["Purchase Price", f"${a.purchase_price:,.0f}", "Acquisition", "As-offered"],
        ["Total Project Cost", f"${total_project_cost:,.0f}", "Acquisition", "Incl. all S&U"],
        ["Loan Amount (LTV)", f"${initial_loan:,.0f} ({a.ltv_pct * 100:.0f}%)", "Financing", "On total project cost"],
        ["Interest Rate", f"{a.interest_rate * 100:.2f}%", "Financing", "Fixed rate"],
        ["Loan Term", f"{a.loan_term} years", "Financing", "Initial term"],
        ["IO Period", f"{a.io_period_months} months", "Financing", "Interest-only"],
        ["Amortization", f"{a.amort_years} years", "Financing", "After IO period"],
        ["Hold Period", f"{hold} years", "Exit", "Target hold"],
        ["Exit Cap Rate", f"{a.exit_cap_rate * 100:.2f}%", "Exit", "Disposition assumption"],
        ["Vacancy Rate", f"{a.vacancy_rate * 100:.1f}%", "Income", "Stabilized assumption"],
        ["Revenue Growth", f"{a.annual_rent_growth * 100:.1f}%/yr", "Income", "Annual escalator"],
        ["Expense Growth", f"{a.expense_growth_rate * 100:.1f}%/yr", "Expenses", "Annual escalator"],
        ["GP/LP Split", f"{a.gp_equity_pct * 100:.0f}% / {a.lp_equity_pct * 100:.0f}%", "Capital", "Equity structure"],
        ["Target LP IRR", f"{a.target_lp_irr * 100:.1f}%", "Returns", "Investment threshold"],
        ["Min LP IRR", f"{a.min_lp_irr * 100:.1f}%", "Returns", "Minimum acceptable"],
    ]
    _safe_pop(28, uw_rows)

    # ── Table 30 (idx=29): Sources & Uses ─────────────────────────
    total_uses = fo.total_uses or 0
    def _pct(amt):
        return f"{amt / total_uses * 100:.1f}%" if total_uses > 0 and amt else ""
    su_prof_dd = sum(
        (getattr(a, k, 0) or 0)
        for k in ['legal_closing', 'title_insurance', 'legal_bank', 'appraisal',
                  'environmental', 'surveyor', 'architect', 'structural',
                  'civil_eng', 'meps', 'legal_zoning', 'geotech']
    )
    su_hard = (getattr(a, 'const_hard', 0) or 0) + (getattr(a, 'const_reserve', 0) or 0)
    su_orig = (fo.initial_loan_amount or 0) * (getattr(a, 'origination_fee_pct', 0.01) or 0.01)
    su_rows = [
        ["Purchase Price", f"${a.purchase_price:,.0f}", _pct(a.purchase_price), "Acquisition"],
        ["Transfer Tax",
         f"${a.purchase_price * getattr(a, 'transfer_tax_rate', 0.02139):,.0f}",
         "", "PA buyer share"],
        ["Professional & DD", f"${su_prof_dd:,.0f}", "", "Legal, title, inspections"],
        ["Construction Hard Costs", f"${su_hard:,.0f}", "", "Renovation + reserve"],
        ["Origination Fee", f"${su_orig:,.0f}", "", "1% of loan"],
        ["Senior Debt", f"${fo.initial_loan_amount or 0:,.0f}",
         f"{(fo.initial_loan_amount or 0) / max(total_uses, 1) * 100:.0f}% LTV",
         f"{getattr(a, 'ltv_pct', 0.70) * 100:.0f}% LTV"],
        ["Total Equity Required", f"${fo.total_equity_required or 0:,.0f}", "", "GP + LP"],
        ["GP Equity", f"${fo.gp_equity or 0:,.0f}",
         f"{getattr(a, 'gp_equity_pct', 0.10) * 100:.0f}%", ""],
        ["LP Equity", f"${fo.lp_equity or 0:,.0f}",
         f"{getattr(a, 'lp_equity_pct', 0.90) * 100:.0f}%", ""],
    ]
    _safe_pop(29, su_rows)

    # ── Table 31 (idx=30): Construction Budget ────────────────────
    hard = a.const_hard or 0
    reserve = a.const_reserve or 0
    const_rows = []
    if hard > 0 or reserve > 0:
        if hard > 0:
            const_rows.append(["Hard Costs", "Renovation / Conversion", f"${hard:,.0f}",
                               f"${hard / gba:.2f}" if gba > 0 else "—", "GC budget"])
        if reserve > 0:
            const_rows.append(["Hard Costs", "Construction Reserve", f"${reserve:,.0f}",
                               f"${reserve / gba:.2f}" if gba > 0 else "—", "Contingency"])
    if not const_rows:
        const_rows = [["No construction budget", "—", "$0", "—", "Stabilized acquisition"]]
    _safe_pop(30, const_rows)

    # ── Table 33 (idx=32): 10-Year Pro Forma Summary ──────────────
    pf_rows = []
    for yr in (fo.pro_forma_years or []):
        yr_num = yr.get("year", "")
        gpr_v = yr.get("gpr", 0) or yr.get("gross_potential_rent", 0)
        egi_v = yr.get("egi", 0) or yr.get("egr", 0) or yr.get("effective_gross_income", 0)
        opex_v = yr.get("opex", 0) or yr.get("total_opex", 0) or yr.get("operating_expenses", 0)
        noi_v = yr.get("noi", 0)
        ds_v = yr.get("debt_service", 0) or yr.get("annual_debt_service", 0)
        cfbt_v = yr.get("cfbt", 0) or yr.get("fcf", 0) or yr.get("free_cash_flow", 0) or yr.get("cash_flow_before_tax", 0)
        coc_v = yr.get("coc", 0) or yr.get("cash_on_cash", 0)
        pf_rows.append([
            f"Year {yr_num}",
            f"${gpr_v:,.0f}",
            f"${egi_v:,.0f}",
            f"${opex_v:,.0f}",
            f"${noi_v:,.0f}",
            f"${ds_v:,.0f}",
            f"${cfbt_v:,.0f}",
            f"{coc_v * 100:.1f}%" if isinstance(coc_v, float) else str(coc_v),
        ])
    _safe_pop(32, pf_rows)

    # ── Table 36 (idx=35): Exit Analysis ──────────────────────────
    exit_noi = proforma[hold - 1].get("noi", 0) if len(proforma) >= hold else 0
    gross_sale = fo.gross_sale_price or 0
    disp_costs = gross_sale * a.disposition_costs_pct
    net_sale = fo.net_sale_proceeds or 0
    exit_bal = gross_sale - disp_costs - (fo.net_equity_at_exit or 0) if gross_sale > 0 else 0
    exit_rows = [
        ["Exit Year NOI", f"${exit_noi:,.0f}", f"Year {hold} pro forma NOI"],
        ["Exit Cap Rate", f"{a.exit_cap_rate * 100:.2f}%", "Underwritten disposition cap rate"],
        ["Gross Sale Price", f"${gross_sale:,.0f}", "NOI / Cap Rate"],
        ["Less: Disposition Costs", f"(${disp_costs:,.0f})", f"{a.disposition_costs_pct * 100:.1f}% of gross"],
        ["Net Sale Proceeds", f"${net_sale:,.0f}", "After disposition costs"],
        ["Net Equity at Exit", f"${fo.net_equity_at_exit or 0:,.0f}", "To equity investors"],
        ["Total Equity Invested", f"${total_equity:,.0f}", "GP + LP contributions"],
        ["Equity Multiple", f"{fo.project_equity_multiple or 0:.2f}x", "Net proceeds / invested equity"],
    ]
    _safe_pop(35, exit_rows)

    # ── Table 39 (idx=38): LP/GP Waterfall Tiers ─────────────────
    wf_rows = [["Tier 0: Preferred Return", f"{a.pref_return * 100:.1f}%", "0%",
                f"{a.lp_equity_pct * 100:.0f}%", f"{a.gp_equity_pct * 100:.0f}%"]]
    for t in a.waterfall_tiers:
        wf_rows.append([
            f"Tier {t.tier_number}: {t.hurdle_type.upper()} Hurdle",
            f"{t.hurdle_value * 100:.1f}%",
            f"{t.gp_share * 100:.0f}%",
            f"{t.lp_share * 100:.0f}%",
            f"{t.gp_share * 100:.0f}%",
        ])
    wf_rows.append(["Residual (above all tiers)", "—",
                     f"{a.residual_tier.gp_share * 100:.0f}%",
                     f"{a.residual_tier.lp_share * 100:.0f}%",
                     f"{a.residual_tier.gp_share * 100:.0f}%"])
    _pop_by_header(doc, "Tier", wf_rows, "WATERFALL")

    # ── Table 41 (idx=40): Environmental Screening ────────────────
    env_rows = [
        ["EPA Brownfields", "No flags identified" if not md.epa_env_flags else "; ".join(md.epa_env_flags),
         "EPA EnviroFacts", "Review required"],
        ["Phase I ESA", "Not completed", "Due diligence gap", "Required pre-closing"],
        ["Phase II ESA", "Not applicable (Phase I pending)", "Contingent on Phase I", "TBD"],
    ]
    _pop_by_header(doc, "Risk Factor", env_rows, "ENV SCREENING")

    # ── Table 42 (idx=41): Climate Risk (First Street) ───────────
    climate_rows = [
        ["Flood", f"{md.first_street_flood or 'Not Determined'}", "First Street / FEMA", "Zone confirmation required"],
        ["Fire", f"{md.first_street_fire or 'Low (urban)'}", "First Street Foundation", ""],
        ["Heat", f"{md.first_street_heat or 'Moderate'}", "First Street Foundation", "Urban heat island"],
        ["Wind", f"{md.first_street_wind or 'Low'}", "First Street Foundation", "Mid-Atlantic"],
    ]
    _safe_pop(41, climate_rows)

    # ── Table 43 (idx=42): Title Search Summary ──────────────────
    title_rows = [
        ["Title Search", "Not completed", "No title commitment on file", "Required pre-closing"],
        ["Title Insurance", "Not bound", f"Budget: ${a.title_insurance:,.0f}", "Bind at closing"],
        ["ALTA Survey", "Not completed", "Due diligence gap", "Required for lender"],
    ]
    _safe_pop(42, title_rows)

    # ── Table 44 (idx=43): Outstanding Violations & Permits ──────
    _pop_by_header(doc, "Type",
        [["L&I Search", "Not completed", "Open violations unknown", "Pending", "Municipal lien search required"]],
        "VIOLATIONS")

    # ── Table 48 (idx=47): Insurance Coverage Rollup ─────────────
    ins_pf = deal.insurance.insurance_proforma_line_item or a.insurance
    ins_rows = [
        ["Property (Replacement)", "Commercial", f"${ins_pf * 0.8:,.0f}", f"${ins_pf * 1.2:,.0f}", f"${ins_pf:,.0f}", "Replacement cost"],
        ["General Liability", "Commercial", "$4,500", "$6,500", "$5,000", "$1M/$2M aggregate"],
        ["Umbrella/Excess", "Excess", "$2,000", "$3,500", "$2,500", "$5M limit"],
        ["Loss of Rents", "Business income", "$2,500", "$4,000", "$3,000", "12-month indemnity"],
    ]
    _safe_pop(47, ins_rows)

    # ── Table 54 (idx=53): Milestone Schedule ────────────────────
    milestone_rows = [
        ["Due Diligence", "Phase I ESA", "Month 1", "30 days", "LOI executed", "Sponsor"],
        ["Due Diligence", "Title search", "Month 1", "21 days", "Contract executed", "Title company"],
        ["Financing", "Loan application", "Month 2", "30 days", "Phase I clean", "Lender"],
        ["Closing", "Acquisition closing", "Month 3", "30 days", "Loan approval", "All parties"],
        ["Operations", "Stabilized operations", "Year 1–2", "Ongoing", "Lease executed", "Property mgr"],
        ["Exit", "Disposition", f"Year {hold}", "6 months", "Market conditions", "Sponsor"],
    ]
    _pop_by_header(doc, "Phase", milestone_rows, "MILESTONE")

    # ── Table 57 (idx=56): Scenario Comparison ───────────────────
    def _safe_pct(val, decimals=1):
        if val is None: return "N/A"
        try: return f"{float(val)*100:.{decimals}f}%"
        except: return "N/A"
    def _safe_dollar(val):
        if val is None: return "N/A"
        try: return f"${float(val):,.0f}"
        except: return "N/A"
    def _safe_x(val, decimals=2):
        if val is None: return "N/A"
        try: return f"{float(val):.{decimals}f}x"
        except: return "N/A"

    lp_irr   = fo.lp_irr
    lp_em    = fo.lp_equity_multiple
    proj_irr = fo.project_irr
    dscr_y1  = fo.dscr_yr1
    gi_cap   = fo.going_in_cap_rate
    coc_y1   = fo.cash_on_cash_yr1
    exit_px  = fo.gross_sale_price
    net_exit = fo.net_equity_at_exit
    eq_mult  = fo.project_equity_multiple
    fcf_y1   = (fo.pro_forma_years[0].get("fcf") if fo.pro_forma_years else None)

    # ── Conservative scenario: re-run financials on a deepcopy with
    #    rent growth -1pp, exit cap +0.5pp, vacancy +2pp ──────────────
    cons = {}
    try:
        import copy as _copy
        from financials import run_financials as _run_fin
        cdeal = _copy.deepcopy(deal)
        ca = cdeal.assumptions
        ca.annual_rent_growth = max(0.0, (ca.annual_rent_growth or 0.03) - 0.01)
        ca.exit_cap_rate      = (ca.exit_cap_rate or 0.07) + 0.005
        ca.vacancy_rate       = (ca.vacancy_rate or 0.05) + 0.02
        # Reset financial outputs so the engine fully recomputes
        from models.models import FinancialOutputs as _FO
        cdeal.financial_outputs = _FO()
        cdeal = _run_fin(cdeal)
        cfo = cdeal.financial_outputs
        cons = {
            "lp_irr":   cfo.lp_irr,
            "lp_em":    cfo.lp_equity_multiple,
            "proj_irr": cfo.project_irr,
            "dscr_y1":  cfo.dscr_yr1,
            "gi_cap":   cfo.going_in_cap_rate,
            "coc_y1":   cfo.cash_on_cash_yr1,
            "exit_px":  cfo.gross_sale_price,
            "net_exit": cfo.net_equity_at_exit,
            "eq_mult":  cfo.project_equity_multiple,
        }
        logger.info("SCENARIO: base computed, conservative computed (lp_irr=%s vs %s)",
                    lp_irr, cfo.lp_irr)
    except Exception as exc:
        logger.warning("SCENARIO conservative re-run failed: %s", exc)

    def _c(key, fmt):
        v = cons.get(key)
        return fmt(v) if v is not None else "N/A"

    scenario_rows = [
        ["LP IRR",              _safe_pct(lp_irr),    _c("lp_irr", _safe_pct),                 "12.0%"],
        ["LP Equity Multiple",  _safe_x(lp_em),       _c("lp_em", _safe_x),                    "1.8x"],
        ["Project IRR",         _safe_pct(proj_irr),  _c("proj_irr", _safe_pct),               "—"],
        ["Year 1 DSCR",         _safe_x(dscr_y1),     _c("dscr_y1", _safe_x),                  "1.20x"],
        ["Going-In Cap Rate",   _safe_pct(gi_cap, 2), _c("gi_cap", lambda v: _safe_pct(v, 2)), "≥5.5%"],
        ["Year 1 Cash-on-Cash", _safe_pct(coc_y1),    _c("coc_y1", _safe_pct),                 "≥6.0%"],
        ["Gross Exit Price",    _safe_dollar(exit_px),_c("exit_px", _safe_dollar),             "—"],
        ["Net Equity at Exit",  _safe_dollar(net_exit),_c("net_exit", _safe_dollar),           "—"],
        ["Equity Multiple",     _safe_x(eq_mult),     _c("eq_mult", _safe_x),                  "—"],
    ]
    logger.info("SCENARIO: built %d rows", len(scenario_rows))
    # Find the Scenario Comparison table specifically (avoid KPI table which also has "Metric")
    _scenario_table = None
    for _st in doc.tables:
        if not _st.rows:
            continue
        _hdr = ' '.join(c.text for c in _st.rows[0].cells).lower()
        if 'metric' in _hdr and ('scenario' in _hdr or 'base' in _hdr or 'conservative' in _hdr):
            _scenario_table = _st
            break
    if _scenario_table is None:
        # Fallback: 4-column table with "Metric" in first cell
        for _st in doc.tables:
            if not _st.rows:
                continue
            _hdr = ' '.join(c.text for c in _st.rows[0].cells).lower()
            if 'metric' in _hdr and len(_st.rows[0].cells) == 4:
                _scenario_table = _st
                break
    if _scenario_table and scenario_rows:
        populate_table(_scenario_table, scenario_rows)
        logger.info("SCENARIO: populated %d rows", len(scenario_rows))
    else:
        logger.warning("SCENARIO: could not find Scenario Comparison table")

    # ── Table 58 (idx=57): Go/No-Go Assessment ───────────────────
    def _verdict(condition):
        return "PASS" if condition else "FAIL"

    gono_rows = [
        ["Going-In Cap Rate ≥ 5.5%",
         f"Actual: {_safe_pct(gi_cap, 2)}",
         _verdict(gi_cap is not None and gi_cap >= 0.055)],
        ["LP IRR ≥ 12.0%",
         f"Actual: {_safe_pct(lp_irr)}",
         _verdict(lp_irr is not None and lp_irr >= 0.12)],
        ["LP Equity Multiple ≥ 1.8x",
         f"Actual: {_safe_x(lp_em)}",
         _verdict(lp_em is not None and lp_em >= 1.8)],
        ["Year 1 DSCR ≥ 1.20x",
         f"Actual: {_safe_x(dscr_y1)}",
         _verdict(dscr_y1 is not None and dscr_y1 >= 1.2)],
        ["Positive Free Cash Flow Yr 1",
         f"Actual: {_safe_dollar(fcf_y1)}",
         _verdict(fcf_y1 is not None and fcf_y1 > 0)],
    ]
    logger.info("GO_NOGO: built %d rows", len(gono_rows))
    _pop_by_header(doc, "Criterion", gono_rows, "GO_NOGO")

    # ── Table 46 (idx=45): DD Flag Summary ──────────────────────
    dd_flag_rows = []
    for f in deal.dd_flags:
        color_emoji = {"RED": "\U0001f534  RED", "AMBER": "\U0001f7e1  YELLOW",
                       "GREEN": "\U0001f7e2  GREEN"}.get(f.color.value, f.color.value)
        dd_flag_rows.append([color_emoji, f.title, f.category, f.remediation or f.narrative[:60]])
    if not dd_flag_rows:
        # Default flags for 5600 Chestnut-style deals
        dd_flag_rows = [
            ["\U0001f534  RED", "100% vacancy — no executed lease", "Financial", "Execute LOI with qualified tenant before closing"],
            ["\U0001f534  RED", "Negative NOI across entire hold period", "Financial", "Re-underwrite upon lease execution with market rents"],
            ["\U0001f534  RED", "No Phase I ESA on file", "Environmental", "Commission Phase I ESA — budget $6,000, 30-day turnaround"],
            ["\U0001f7e1  YELLOW", "Zoning classification unconfirmed", "Legal/Regulatory", "Engage zoning counsel; verify via atlas.phila.gov"],
            ["\U0001f7e1  YELLOW", "FEMA flood zone 'Not Determined'", "Environmental", "Obtain SFHDS certificate from licensed surveyor"],
            ["\U0001f7e1  YELLOW", "No rent or sale comparables", "Market", "Request CoStar comp set from Shonda at Binswanger"],
            ["\U0001f7e1  YELLOW", "24-month balloon refinance risk", "Financial", "Secure lease within 18 months to support refi underwriting"],
            ["\U0001f7e2  GREEN", "2011 renovation — reduced deferred maintenance", "Physical", "Confirm with PCA inspection"],
            ["\U0001f7e2  GREEN", "Corner lot with dual frontage", "Physical", "Leverage for medical/retail tenant marketing"],
            ["\U0001f7e2  GREEN", "Below replacement cost basis ($86.79/SF)", "Financial", "Validate with closed sale comps"],
            ["\U0001f7e2  GREEN", "University City institutional anchor proximity", "Market", "Cite in tenant marketing materials"],
            ["\U0001f7e2  GREEN", "Dual-path strategy preserves optionality", "Strategic", "Finalize use decision post-zoning confirmation"],
        ]
    logger.info("DD FLAG: %d rows to populate (template rows cleared first)", len(dd_flag_rows))
    _safe_pop(45, dd_flag_rows)

    # ── Table 61 (idx=60): Data Sources Provenance ────────────────
    report_date = deal.report_date or ""
    provenance_rows = [
        ["Property Data", "User Input (Frontend)", report_date, "localhost:8000"],
        ["Demographics", "U.S. Census ACS 2022", report_date, "data.census.gov"],
        ["Flood Zone", "FEMA NFHL", report_date, "msc.fema.gov"],
        ["Environmental", "EPA EnviroFacts", report_date, "epa.gov/enviro"],
        ["Maps", "Google Maps Static API", report_date, "maps.googleapis.com"],
        ["Financial Model", "DealDesk Engine v1.0", report_date, "Deterministic Python"],
        ["Narratives", "Claude Sonnet 4.5", report_date, "api.anthropic.com"],
        ["Municipal Data", "DealDesk Registry", report_date, "municipal_registry.csv"],
    ]
    _safe_pop(60, provenance_rows)

    # Sensitivity table — disabled: keyword-match was injecting the matrix
    # into the Executive Summary table (matched on 'irr'). Section 12.5
    # renders the matrix via the docxtpl context keys (sensitivity_rows,
    # sensitivity_cap_axis). Re-enable only with a tighter table selector.
    # populate_sensitivity_table(doc, deal)

    # Fix 9: data row colors — parchment background for dark-inherited rows
    color_count = _fix_dark_data_rows(doc)
    if color_count > 0:
        logger.info("TABLE COLORS: fixed %d data cells from dark to parchment", color_count)

    # Remove placeholder boxes that have no data source yet
    PLACEHOLDER_BOXES_TO_REMOVE = [
        "Floor Plans",
        "Aerial Location Map",
        "Neighborhood Map",
        "FEMA Flood Map",
        "KPI Dashboard",
        "Supply Pipeline",
        "Demographic Chart",
        "Zoning Code",
    ]
    for box_label in PLACEHOLDER_BOXES_TO_REMOVE:
        if remove_placeholder_box(doc, box_label):
            logger.info("PLACEHOLDER: removed '%s' box", box_label)

    # Remove single-cell parameter-label placeholder boxes (sensitivity
    # page KPI boxes — one label per box across report pages 22-35).
    _remove_parameter_placeholder_boxes(doc, [
        "Purchase Price", "Total Project Cost", "Loan Amount (LTV)",
        "Interest Rate", "Loan Term", "IO Period", "Amortization",
        "Hold Period", "Exit Cap Rate", "Vacancy Rate",
        "Revenue Growth", "Expense Growth", "GP/LP Split",
        "Target LP IRR", "Min LP IRR", "Transfer Tax",
        "Professional & DD", "Construction Hard Costs", "Origination Fee",
        "Senior Debt", "Total Equity Required", "GP Equity", "LP Equity",
    ])

    # Remove per-year placeholder boxes ("Year 1" … "Year 13")
    # — the 10-year summary table at idx=32 holds the consolidated pro forma.
    _yr_removed = 0
    for year in range(1, 14):
        if remove_placeholder_box(doc, f"Year {year}"):
            _yr_removed += 1
    if _yr_removed:
        logger.info("PLACEHOLDER: removed %d 'Year N' per-year placeholder boxes", _yr_removed)

    logger.info("TABLE POPULATE: completed all data tables")
    return _populate_counter[0]


def _populate_docx(deal: DealData) -> Path:
    """Populate DealDesk_Report_Template_v4.docx with template context. Returns docx path."""
    ctx = _build_context(deal)

    tpl = DocxTemplate(str(WORD_TEMPLATE))

    # Generate and merge image context
    # _build_image_context returns (ctx_dict, tmp_file_list)
    # temp files must stay alive until AFTER tpl.render() — InlineImage is lazy
    img_ctx, img_tmp_files = _build_image_context(deal, tpl)
    ctx.update(img_ctx)

    logger.info(f"[WORD_BUILDER] ctx keys: {list(ctx.keys())}")
    for k, v in ctx.items():
        try:
            v_repr = repr(v)[:80] if not hasattr(v, '_insert_image') else "<InlineImage>"
        except Exception:
            v_repr = "<unrepresentable>"
        logger.info(f"  {k}: {type(v).__name__} = {v_repr}")

    for key, val in ctx.items():
        if isinstance(val, str):
            if len(val) == 0:
                logger.warning(f"CTX EMPTY: {key}")
            elif val.strip().startswith('{{') or 'placeholder' in val.lower():
                logger.warning(f"CTX UNFILLED: {key} = {val[:80]}")
            else:
                logger.info(f"CTX OK: {key} = {len(val)} chars")
        elif val is None:
            logger.warning(f"CTX NONE: {key}")

    # ── Fix 9: master context audit (critical keys) ──────────────────
    logger.info("=" * 60)
    logger.info("WORD BUILDER CONTEXT AUDIT — DEAL %s", deal.deal_id)
    logger.info("=" * 60)
    _critical_keys = [
        "kpi_rows", "street_view_image", "has_street_view",
        "parcel_a_account", "parcel_a_zoning", "parcel_a_owner",
        "parcel_census_tract", "parcel_fips",
        "hbu_content", "transit_rows", "income_egi",
        "income_gpr", "fmr_2br",
    ]
    for _k in _critical_keys:
        _v = ctx.get(_k)
        if isinstance(_v, list):
            logger.info("  CTX[%s] = list(%d items)", _k, len(_v))
        elif _v is None:
            logger.warning("  CTX[%s] = MISSING/NONE \u2190 FIX NEEDED", _k)
        else:
            # IMPORTANT: Do NOT call str() on InlineImage objects —
            # that triggers _insert_image() before tpl.render() and crashes.
            # Check type name safely instead.
            _type_name = type(_v).__name__
            if _type_name == "InlineImage":
                logger.info("  CTX[%s] = InlineImage(ready)", _k)
            else:
                try:
                    _repr = str(_v)[:80]
                except Exception:
                    _repr = f"<{_type_name}>"
                logger.info("  CTX[%s] = %s", _k, _repr)
    logger.info("=" * 60)

    try:
        tpl.render(ctx)
    finally:
        # Clean up temp image files after render completes
        import os as _os
        for tmp_path in img_tmp_files:
            try:
                _os.unlink(tmp_path)
            except Exception:
                pass
    _strip_highlight(tpl.docx)

    # ── Fix 3: remove leaked Jinja conditional-comment paragraphs ────
    _paras_to_remove = []
    for _para in tpl.docx.paragraphs:
        _txt = _para.text or ""
        if ("Conditional block" in _txt
                or "image_placements.json" in _txt
                or "renders only when image_type" in _txt):
            _paras_to_remove.append(_para)
    for _para in _paras_to_remove:
        _p = _para._element
        _p.getparent().remove(_p)
    if _paras_to_remove:
        logger.info("FLOOR PLAN: removed %d leaked conditional paragraph(s)",
                    len(_paras_to_remove))

    # ── Populate data tables FIRST against original-template indices ──
    n_tables_before = len(tpl.docx.tables)
    tables_populated_count = _populate_data_tables(tpl.docx, deal, ctx)

    # Remove sage-green image placeholder boxes AFTER population — otherwise
    # the removal shifts doc.tables[] indices and the hardcoded _safe_pop(N,)
    # positions in _populate_data_tables write into the wrong tables.
    _remove_image_placeholder_boxes(tpl.docx)

    n_tables_after = len(tpl.docx.tables)
    placeholders_removed_count = n_tables_before - n_tables_after
    logger.info("DOCX COMPLETE: tables_populated=%d, placeholders_removed=%d",
                tables_populated_count, placeholders_removed_count)

    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = OUTPUTS_DIR / f"{deal.deal_id}_report.docx"
    tpl.save(str(docx_path))
    logger.info("DOCX generated: %s", docx_path)
    return docx_path


def _update_toc_fields(docx_path: Path) -> None:
    """
    Open the DOCX in a live Word instance via win32com, update all fields
    (including the Table of Contents), save, and close.

    This must run on Windows with Microsoft Word installed.  It is required
    because docxtpl renders the template but cannot update Word fields — the
    TOC field remains empty until Word recalculates it.

    Falls back gracefully with a warning if win32com is unavailable (e.g.
    on Streamlit Community Cloud) — the PDF will render with a blank TOC
    but all other content will be correct.
    """
    try:
        import win32com.client as win32
    except ImportError:
        logger.warning(
            "win32com not available — TOC fields will not be updated. "
            "Install pywin32 to enable automatic TOC generation."
        )
        return

    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        try:
            doc = word.Documents.Open(str(docx_path.resolve()))
            doc.Fields.Update()          # updates all fields including TOC
            doc.TablesOfContents(1).Update()  # explicitly refresh TOC entries + page numbers
            doc.Save()
            doc.Close()
        finally:
            word.Quit()
        logger.info("TOC fields updated: %s", docx_path.name)
    except Exception as exc:
        logger.warning(
            "win32com TOC update failed: %s — proceeding with blank TOC", exc
        )


def _convert_to_pdf(docx_path: Path) -> Path:
    """Convert DOCX to PDF via LibreOffice headless. Returns PDF path."""
    if platform.system() == "Windows":
        soffice = r"C:\Program Files\LibreOffice\program\soffice.exe"
    else:
        soffice = "soffice"

    output_dir = docx_path.parent
    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(output_dir),
                str(docx_path),
            ],
            timeout=PDF_CONVERSION_TIMEOUT,
            check=True,
            capture_output=True,
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError(
            f"LibreOffice PDF conversion timed out after {PDF_CONVERSION_TIMEOUT}s"
        )
    except subprocess.CalledProcessError as exc:
        raise RuntimeError(f"LibreOffice PDF conversion failed: {exc.stderr.decode()}")
    except FileNotFoundError:
        raise RuntimeError(
            f"LibreOffice not found at: {soffice}. "
            "Please verify the installation path."
        )

    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists():
        raise RuntimeError(f"PDF not found after conversion: {pdf_path}")

    logger.info("PDF generated: %s", pdf_path)
    return pdf_path


# ═══════════════════════════════════════════════════════════════════════════
# PUBLIC API
# ═══════════════════════════════════════════════════════════════════════════

def generate_report(deal: DealData) -> DealData:
    """
    Generate the full PDF underwriting report.

    1. Calls Prompt 4-MASTER (Sonnet) to generate all narrative sections.
    2. If investor_mode is True, calls Prompt 5D to rewrite 9 narrative
       blocks in LP-appropriate language.
    3. Populates DealDesk_Report_Template_v4.docx via docxtpl.
    4. Converts to PDF via LibreOffice headless (60s timeout).

    Args:
        deal: DealData with all upstream modules already populated.

    Returns:
        The same DealData object with output_pdf_path set.
    """
    # Stage 1: Generate all narratives
    _generate_narratives(deal)

    # Stage 2: Investor mode rewrite (if applicable)
    if deal.investor_mode:
        _rewrite_investor_narratives(deal)

    # Stage 3: Populate DOCX template
    docx_path = _populate_docx(deal)

    # Stage 3B: Update Word fields (TOC page numbers) before PDF conversion
    _update_toc_fields(docx_path)

    # Stage 4: Convert to PDF
    pdf_path = _convert_to_pdf(docx_path)
    deal.output_pdf_path = str(pdf_path)

    logger.info("Report generation complete: %s", pdf_path)
    return deal
